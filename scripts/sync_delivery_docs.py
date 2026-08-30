"""Sincroniza memoria y guía extensa con la evidencia actual del repositorio."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]


def _set_text(paragraph: Paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _all_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_prefix(doc: DocumentObject, prefix: str, replacement: str) -> int:
    changed = 0
    for paragraph in _all_paragraphs(doc):
        if paragraph.text.startswith(prefix):
            _set_text(paragraph, replacement)
            changed += 1
    return changed


def replace_fragment(doc: DocumentObject, old: str, new: str) -> int:
    changed = 0
    for paragraph in _all_paragraphs(doc):
        if old in paragraph.text:
            _set_text(paragraph, paragraph.text.replace(old, new))
            changed += 1
    return changed


def replace_next_nonempty(doc: DocumentObject, prefix: str, replacement: str) -> int:
    """Sustituye el primer párrafo con contenido posterior a un marcador."""
    for index, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.startswith(prefix):
            continue
        for target in doc.paragraphs[index + 1 :]:
            if target.text.strip():
                _set_text(target, replacement)
                return 1
    return 0


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Inserta un párrafo conservando el estilo del párrafo de referencia."""
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if paragraph.style is not None:
        inserted.style = paragraph.style
    _set_text(inserted, text)
    return inserted


def _keep_table_rows_together(doc: DocumentObject) -> None:
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))


def _add_page_field(paragraph: Paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for element in (begin, instruction, separate, result, end):
        run.append(element)


def _restore_page_fields(doc: DocumentObject) -> None:
    """Restaura PAGE si una sustitución histórica dejó el resultado literal."""

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            marker = "Página 1"
            if marker not in paragraph.text:
                continue
            prefix, suffix = paragraph.text.split(marker, maxsplit=1)
            paragraph.clear()
            paragraph.add_run(prefix + "Página ")
            _add_page_field(paragraph)
            if suffix:
                paragraph.add_run(suffix)


def _replace_controlled_results_table(doc: DocumentObject) -> None:
    for table in doc.tables:
        if not table.rows or table.rows[0].cells[0].text.strip() not in {"Caso", "Modo"}:
            continue
        rows = [
            ["Modo", "N", "Accuracy", "Precisión", "Recall", "F1"],
            ["Heurístico", "16", "100,0 %", "100,0 %", "100,0 %", "100,0 %"],
            ["Neuronal", "16", "75,0 %", "75,0 %", "75,0 %", "75,0 %"],
            ["Combinado 35/65", "16", "93,8 %", "88,9 %", "100,0 %", "94,1 %"],
        ]
        while len(table.rows) > len(rows):
            table._tbl.remove(table.rows[-1]._tr)
        for row_index, values in enumerate(rows):
            for cell, value in zip(table.rows[row_index].cells, values):
                cell.text = value
        return
    raise ValueError("No se encontró la tabla de resultados controlados en la guía extensa.")


def _replace_signal_weights_table(doc: DocumentObject) -> None:
    """Sincroniza la tabla explicativa con las 31 señales del scorer actual."""

    rows = [
        ["Señal", "Peso", "Señal", "Peso"],
        ["Reply-To diferente", "14", "Nombre visible engañoso", "14"],
        ["Marca engañosa", "10", "Cabecera de spoofing", "10"],
        ["Enlaces sospechosos", "10", "Dominio en blacklist local", "10"],
        ["Autenticación fallida", "14", "DMARC fallido", "8"],
        ["DKIM mal formado", "5", "Received sospechoso", "8"],
        ["Saludo genérico", "6", "Solicitud de credenciales", "6"],
        ["Cambio de datos bancarios", "8", "Transferencia urgente", "8"],
        ["Suplantación ejecutiva", "8", "Message-ID sospechoso", "6"],
        ["Meta refresh", "6", "Redirección JavaScript", "6"],
        ["HTML sospechoso", "6", "Adjunto sospechoso", "6"],
        ["Lenguaje urgente", "6", "Asunto sospechoso", "5"],
        ["Parámetros URL sospechosos", "5", "Punycode/Unicode", "5"],
        ["Incoherencia remitente", "8", "Acortador", "5"],
        ["Texto de enlace distinto", "5", "Formulario HTML", "4"],
        ["Action sospechoso", "5", "Referencia a archivo", "4"],
        ["Firmado/cifrado", "-3", "", ""],
    ]
    for table in doc.tables:
        if not table.rows or [cell.text.strip() for cell in table.rows[0].cells] != rows[0]:
            continue
        while len(table.rows) < len(rows):
            table.add_row()
        while len(table.rows) > len(rows):
            table._tbl.remove(table.rows[-1]._tr)
        for row_index, values in enumerate(rows):
            for cell, value in zip(table.rows[row_index].cells, values):
                cell.text = value
        return
    raise ValueError("No se encontró la tabla de pesos heurísticos en la guía extensa.")


def sync_memory() -> None:
    path = ROOT / "TFG.docx"
    doc = Document(path)
    replacements = [
        (
            "La solución adopta una arquitectura",
            "La solución adopta una arquitectura cliente-servidor. El navegador se comunica con Streamlit, que actúa como capa de presentación y cliente HTTP del backend central; la extensión y el monitor consumen esa misma API. Solo backend_server.py normaliza, analiza, entrena y mantiene los modelos activos. En la configuración académica ambos lados se ejecutan en el mismo equipo y sobre loopback, pero son procesos y responsabilidades independientes.",
        ),
        (
            "La validación del sistema",
            "La validación del sistema se realiza en seis niveles complementarios: 87 pruebas Python de componentes e integración, 2 recorridos reales con Chromium, un benchmark reproducible, una calibración bilingüe de 40 casos, 16 EML locales reservados y un diagnóstico de 1.528 textos externos DIFrauD. Un recorrido levanta cliente web y backend separados; otro prueba Gmail, HTTP y Telegram de extremo a extremo con dobles externos. Los EML son sintéticos y DIFrauD conserva riesgo de solapamiento con fuentes de entrenamiento; ninguna cifra se extrapola a producción.",
        ),
        (
            "La metodología seguida ha combinado",
            "La metodología seguida ha combinado una revisión bibliográfica de fuentes académicas y de industria con un desarrollo iterativo del prototipo. El sistema se ha implementado en Python con una arquitectura cliente-servidor: el navegador usa Streamlit como presentación, mientras Streamlit, la extensión y el monitor envían las entradas por HTTP al backend central. Solo el servidor normaliza, analiza, entrena y mantiene las versiones activas de los modelos. La validación se ha realizado mediante pruebas automatizadas, recorridos funcionales cliente-servidor y evaluación controlada del clasificador en español e inglés.",
        ),
        (
            "Phishing remains one of the most persistent",
            "Phishing remains one of the most persistent and damaging cybersecurity threats. This Bachelor's Thesis designs, implements and evaluates a client-server system for phishing email detection. The browser uses Streamlit as the presentation layer, while Streamlit, the Gmail extension and the monitor send text, EML or structured fields through HTTP to a mandatory central backend. Only the server parses messages, combines explainable heuristics with Spanish and English TF-IDF + MLP models, trains them and keeps one active version per language, so an update is shared by every client. The prototype is validated with automated tests, real browser-to-backend journeys, separate calibration and a reserved local EML scenario corpus, while its limits for production are stated explicitly.",
        ),
        (
            "Una vez consolidado el marco teórico",
            "Una vez consolidado el marco teórico, se adoptó un enfoque de desarrollo en espiral adaptado al tamaño del trabajo, llevado a cabo entre febrero y junio de 2026. Primero se construyó un núcleo heurístico explicable, después se añadió el clasificador neuronal TF-IDF + MLP y, a continuación, se separaron el backend HTTP central y los clientes Streamlit, extensión y monitor. Gmail aporta mensajes autorizados y Telegram actúa como salida opcional; ambos se coordinan alrededor del contrato del servidor. Cada avance se cerró con la ejecución de la suite de pruebas, manteniendo el prototipo en un estado funcional durante todo el proceso.",
        ),
        (
            "La carga de trabajo del segundo cuatrimestre",
            "La carga de trabajo del segundo cuatrimestre se concentró en el motor heurístico, el clasificador neuronal y el backend central, que reúnen la lógica del prototipo. Después se adaptaron Streamlit, Gmail, la extensión y el monitor como clientes del mismo contrato HTTP. La extensión llama directamente al servidor; el proceso del puerto 8765 se conserva únicamente como proxy de compatibilidad, sin reglas ni modelos propios.",
        ),
        (
            "RiskScorer asigna una ponderación",
            "RiskScorer asigna una ponderación a cada señal activa y calcula una puntuación final entre 0 y 100. Si la puntuación supera el umbral configurado, el mensaje se clasifica como phishing probable. ExplanationBuilder transforma las señales en explicaciones comprensibles. El backend central devuelve ese resultado por HTTP a Streamlit, la extensión o el monitor; los clientes solo lo presentan o, en el caso del monitor, pueden generar una alerta de Telegram y registrar el identificador para evitar duplicados.",
        ),
        (
            "Una vez normalizado el correo",
            "Una vez normalizado el correo, SignalBuilder ejecuta 31 reglas y produce un diccionario estable de señales. Además de cabeceras, autenticación, URLs, HTML y adjuntos, incorpora tres patrones bilingües para BEC sin enlace: cambio de datos bancarios, orden urgente de transferencia y suplantación de un directivo bajo aislamiento o confidencialidad. Si el modo lo requiere, el backend detecta el idioma y reutiliza el modelo neuronal central correspondiente.",
        ),
        (
            "El Trabajo Fin de Grado ha alcanzado",
            "El Trabajo Fin de Grado ha alcanzado los objetivos planteados. Se ha estudiado el phishing desde su base psicológica y su evolución reciente, se ha revisado el estado del arte en machine learning y deep learning y se ha desarrollado un sistema cliente-servidor funcional que combina análisis heurístico explicable con un clasificador neuronal TF-IDF + MLP. Streamlit, la extensión de Gmail y el monitor envían texto, EML o campos estructurados al backend central obligatorio, que procesa la entrada y mantiene los modelos compartidos.",
        ),
        (
            "El quinto objetivo, integrar",
            "El quinto objetivo, integrar la aplicación con Gmail y proporcionar alertas mediante Telegram, se cumple mediante OAuth de solo lectura y el monitor opcional. Los mensajes autorizados se envían al backend central y Telegram recibe únicamente la alerta configurada; ni Streamlit ni el monitor duplican la lógica de análisis.",
        ),
        (
            "Aplicación web local:",
            "Aplicación web cliente-servidor: El navegador usa Streamlit como capa de presentación y Streamlit llama por HTTP al backend central obligatorio. En la configuración académica ambos procesos se ejecutan en el mismo equipo; la extensión y el monitor consumen el mismo contrato.",
        ),
        (
            "La siguiente estructura recoge",
            "La siguiente estructura recoge los componentes principales de la aplicación cliente-servidor, sus clientes y el backend central de análisis:",
        ),
        (
            "Esta separación facilita la evolución del prototipo.",
            "Esta separación facilita la evolución del prototipo. Pueden añadirse reglas o sustituirse el clasificador en el servidor sin redistribuir los clientes. La versión revisada ejecuta correctamente 87 pruebas Python y 2 recorridos reales con Chromium; GitHub Actions repite pruebas, Ruff, calibración, evaluación reproducible, respaldo de defensa y navegación de interfaz.",
        ),
        (
            "La versión revisada ejecuta",
            "La versión revisada ejecuta 87 pruebas unitarias y de integración mediante unittest. Cubren EML, señales BEC, combinación calibrada, selección determinista de idioma, cliente HTTP, entrenamiento central, Gmail, monitor, heurísticas, clasificador, persistencia, API, autenticación administrativa, extensión, evaluación externa, respaldo de defensa y Telegram. Una prueba integral recorre Gmail EML, JSON, HTTP, backend, estado y alerta Telegram. Además, 2 pruebas con Chromium recorren la extensión y un análisis real entre Streamlit y backend. Todas finalizaron correctamente; los avisos de convergencia proceden de iteraciones reducidas del MLP en pruebas rápidas.",
        ),
        (
            "La metodología cuantitativa separa",
            "La metodología cuantitativa separa ajuste y comprobación. Cuarenta casos controlados y cinco particiones estratificadas calibran el modo combinado en 35 % heurístico, 65 % neuronal, umbral 26 y conservación de evidencia a partir de 70. Después se evalúan 16 EML reservados: el heurístico alcanza 100,0 % de accuracy, recall y F1; el combinado obtiene 93,8 % de accuracy, 100,0 % de recall y 94,1 % de F1; el neuronal obtiene 75,0 % en esas métricas. Un diagnóstico DIFrauD de 1.528 textos externos da 90,8 % de accuracy y 96,4 % de recall al combinado, pero se documenta riesgo de solapamiento. Ninguna muestra estima producción.",
        ),
        (
            "La mezcla ponderada original podía diluir una evidencia fuerte",
            "La mezcla ponderada original podía diluir una evidencia fuerte. Tras incorporar señales bilingües de cambio bancario, transferencia urgente y suplantación ejecutiva, la calibración separada asignó 35 % a heurística y 65 % al modelo, fijó el umbral 26 y mantuvo alta confianza 70. Los indicios BEC aislados pesan poco y solo la coincidencia de los tres recibe el refuerzo de alta confianza. Los dos BEC sin enlace reservados pasan a detectarse. Esto resuelve esos escenarios concretos, no todo BEC posible.",
        ),
        (
            "Se realizaron",
            "Se realizaron 87 pruebas Python y 2 recorridos Chromium. Las comprobaciones verifican los tres modos, calibración, idioma determinista, EML serializable, Gmail a backend y Telegram, cliente HTTP, entrenamiento y versión central, caché, señales BEC, API, límites, token administrativo, evaluación externa, respaldo, extensión y Streamlit. Los 16 EML y DIFrauD añaden medición separada con sus límites expresos; no se presentan como evaluación estadística de producción.",
        ),
        (
            "Los resultados confirman el diseño como base funcional",
            "Los resultados confirman el diseño como base funcional: en los 16 EML reservados, el heurístico obtiene 100,0 % de accuracy y recall; el combinado detecta los ocho phishing, incluido BEC sin enlace, con 93,8 % de accuracy y un falso positivo; el neuronal conserva dos falsos negativos y dos falsos positivos. En DIFrauD el combinado logra 90,8 % de accuracy y 96,4 % de recall. La limitación principal pasa a ser la independencia y actualidad de los datos externos, además de la falta de validación real bilingüe y de controles de producción.",
        ),
        (
            "El sexto objetivo, evaluar el comportamiento",
            "El sexto objetivo, evaluar el comportamiento del sistema, se aborda con 87 pruebas Python, 2 recorridos Chromium —incluido web a backend—, benchmark reproducible, calibración separada de 40 casos, 16 EML reservados y 1.528 textos externos con hashes, matriz de confusión y errores por identificador. La estimación independiente sobre correo real reciente sigue delimitada como trabajo pendiente.",
        ),
        (
            "Entre los logros principales destacan",
            "Entre los logros principales destacan la arquitectura cliente-servidor, una versión activa por idioma compartida por todos los clientes, activación atómica, Gmail/Telegram, 87 pruebas Python, 2 pruebas de navegador, CI, calibración separada, detección BEC y evaluación trazable local/externa. La carga diferida reduce históricamente el arranque de heurísticas un 96,6 % y el de la aplicación un 76,2 %. La fusión 35/65 conserva evidencia de alta confianza.",
        ),
        (
            "Como trabajo futuro prioritario se plantea",
            "Como trabajo futuro prioritario se plantea confirmar la calibración sobre datasets reales, recientes e independientes en español e inglés, con licencia, procedencia temporal y deduplicación frente al entrenamiento, y ampliar la detección BEC más allá de los patrones actuales. También se estudiarán ataques adversariales, transformers/LLM, reputación de dominios, análisis dinámico y etiquetado seguro. Un despliegue multiusuario requerirá autenticación, TLS, rate limiting, secretos e aislamiento de sesiones.",
        ),
        (
            "Este anexo resume",
            "Este anexo resume la ejecución cliente-servidor local. Primero se inicia el backend obligatorio en 127.0.0.1:8766 y después Streamlit en 127.0.0.1:8501. La extensión apunta directamente al backend desde Opciones y el monitor usa PHISHING_BACKEND_URL. El puerto 8765 queda como proxy antiguo opcional. Fuera de loopback los clientes exigen HTTPS y --allow-remote exige un token administrativo de al menos 24 caracteres; todavía hacen falta autenticación de inferencia, rate limiting y controles operativos.",
        ),
        (
            "Verificado: 87 pruebas Python",
            "Verificado: 87 pruebas Python y 2 recorridos reales con Chromium, incluido cliente web a backend.",
        ),
        (
            "Verificado: holdout controlado bilingüe",
            "Verificado: calibración controlada de 40 casos y evaluación final de 16 EML locales reservados, con hashes, métricas y errores por caso; pendiente corpus real representativo.",
        ),
    ]
    replacements.extend(
        [
            (
                "El phishing se ha consolidado",
                "El phishing se ha consolidado como una de las ciberamenazas más persistentes y dañinas. Este Trabajo Fin de Grado diseña, implementa y evalúa un sistema cliente-servidor para detectar phishing en correo. Streamlit, la extensión y el monitor envían las entradas a una API central que combina heurística explicable y modelos TF-IDF + MLP en español e inglés. El servidor es el único responsable de parsear, analizar, entrenar y versionar los modelos, de modo que una actualización se aplica a todos los clientes. La memoria revisa el estado del arte y presenta una validación reproducible, delimitando expresamente sus límites frente a producción.",
            ),
            (
                "El prototipo se centra en el análisis",
                "El prototipo se centra en el análisis de mensajes desde varios clientes. El origen puede ser texto, EML, Gmail, la extensión o JSON. El cliente transporta la entrada y el backend central extrae cabeceras, remitente, asunto, cuerpo, HTML, enlaces, anclas y adjuntos para generar el riesgo heurístico, neuronal o combinado. La interfaz recibe un contrato ya calculado y se limita a presentarlo.",
            ),
            (
                "A nivel interno, el paquete sistema_phishing",
                "A nivel interno, backend_client.py implementa el transporte común de los clientes y http_api.py define las rutas. backend_service.py concentra análisis, datasets, evaluación, entrenamiento, metadatos y activación de versiones. Dentro del servidor, analysis_service.py coordina los modos y la caché lingüística; analizador_email.py normaliza EML; signal_builder.py, scorer.py y explanations.py construyen el resultado; modelo_neural.py encapsula TF-IDF + MLP. Las vistas Streamlit no importan almacenamiento ni ejecutan inferencia.",
            ),
            (
                "El sistema permite entrenar modelos en español",
                "El sistema permite una versión activa en español y otra en inglés. La vista de entrenamiento serializa los CSV y los hiperparámetros y los envía al backend, que valida datos, entrena desde cero, guarda el artefacto de forma atómica e invalida la caché. Evaluación y comparación también se ejecutan en el servidor; los clientes reciben únicamente métricas y metadatos.",
            ),
            (
                "El cliente web se desarrolla con Streamlit",
                "El cliente web se desarrolla con Streamlit y organiza cinco vistas: Inicio, Configuración, Detección, Monitor y Entrenamiento. Comparte un sistema visual adaptable con cabecera de marca, navegación compacta, tarjetas de estado y componentes coherentes. Detección ordena fuente, configuración y resultado en tres pasos; Configuración agrupa conexiones, monitor, backend y red neuronal en pestañas. Sus responsabilidades son recoger entradas, llamar mediante BackendClient y presentar la respuesta, sin cargar modelos ni exponer rutas locales completas de credenciales.",
            ),
            (
                "La integración con Gmail se realiza",
                "La integración con Gmail usa OAuth de solo lectura. La vista de detección obtiene EML autorizados y los remite al backend; el proceso monitor_gmail.py hace lo mismo periódicamente y puede notificar por Telegram. Ambos consumen la misma versión central y un fallo temporal queda controlado para reintento.",
            ),
            (
                "Durante el desarrollo se ha realizado una refactorización",
                "Durante el desarrollo se ha realizado una refactorización orientada a SOLID y a una frontera cliente-servidor explícita. La presentación depende de BackendClient; el servidor separa transporte, casos de uso, señales, puntuación, explicaciones, datasets y aprendizaje. Así, los detalles del modelo permanecen fuera de Streamlit, extensión y monitor.",
            ),
            (
                "El sistema es híbrido y local.",
                "El sistema es cliente-servidor. Streamlit, la extensión y el monitor son clientes HTTP; backend_server.py centraliza parser, análisis, entrenamiento y una versión activa por idioma. En la configuración de defensa todos los procesos están en el mismo equipo y loopback, pero el servidor es obligatorio y puede separarse cambiando PHISHING_BACKEND_URL.",
            ),
            (
                "La aplicación ofrece varias entradas:",
                "La aplicación ofrece texto, EML y Gmail desde Streamlit, además del correo visible en la extensión. Todos esos clientes envían la entrada al backend central. /analyze acepta campos JSON o EML Base64 y devuelve riesgo, señales, explicación, idioma y versión para que cada cliente se limite a presentarlos.",
            ),
            (
                "El clasificador neuronal asigna",
                "El clasificador neuronal del backend asigna la probabilidad y selecciona el modelo por idioma. Streamlit recibe conjuntamente la puntuación neuronal y las señales heurísticas; no carga el pipeline ni recalcula la decisión.",
            ),
            (
                "Desde la vista Monitor",
                "Desde la vista Monitor se consultan correos de Gmail y se envían al backend. Si Telegram está configurado, el cliente puede alertar con el remitente, asunto y puntuación devueltos por el servidor.",
            ),
            (
                "Un caso habitual consiste",
                "Un caso habitual consiste en levantar backend y web, cargar un EML o pegar texto. El cliente envía la entrada, el servidor la procesa con la versión central y devuelve señales y riesgo. Para separar físicamente ambos lados se publica el backend detrás de HTTPS y se cambia PHISHING_BACKEND_URL; la latencia dependerá entonces también de la red.",
            ),
            (
                "Verificado: Arranque y navegación",
                "Verificado: arranque separado de backend y Streamlit, navegación y análisis HTTP completo.",
            ),
        ]
    )
    missing = []
    for prefix, replacement in replacements:
        if not replace_prefix(doc, prefix, replacement) and not any(
            paragraph.text == replacement for paragraph in _all_paragraphs(doc)
        ):
            missing.append(prefix)
    replace_fragment(
        doc,
        "python -m pip install -r requirements.txt",
        "python -m pip install -r requirements-dev.txt -c constraints.txt",
    )
    replace_fragment(
        doc,
        "pip install -r requirements.txt",
        "python -m pip install -r requirements-dev.txt -c constraints.txt",
    )
    replace_fragment(doc, "54 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "59 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "60 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "72 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "81 pruebas Python", "87 pruebas Python")
    replace_fragment(
        doc,
        "API: Interfaz de operaciones que permite a un cliente local enviar datos",
        "API: Interfaz de operaciones que permite a un cliente enviar datos",
    )
    replace_fragment(doc, "# API HTTP local", "# Servidor HTTP central")
    replace_fragment(
        doc,
        "# Configuración de Gmail, Telegram y modelos",
        "# Configuración de clientes y servidor",
    )
    replace_fragment(
        doc,
        "# Detección desde texto, EML y Gmail",
        "# Cliente web de detección",
    )
    replace_fragment(
        doc,
        "# Adaptador de la extensión Gmail",
        "# Proxy legado opcional (puerto 8765)",
    )
    replace_fragment(
        doc,
        "# Coordinación del análisis local",
        "# Coordinación del análisis en servidor",
    )
    replace_fragment(
        doc,
        "# Compatibilidad con el backend",
        "# Casos de uso del servidor central",
    )
    if not any(
        "backend_client.py" in paragraph.text and "Cliente HTTP compartido" in paragraph.text
        for paragraph in doc.paragraphs
    ):
        for paragraph in doc.paragraphs:
            if "backend_service.py" in paragraph.text and "Casos de uso" in paragraph.text:
                insert_paragraph_after(
                    paragraph,
                    "│       ├── backend_client.py         # Cliente HTTP compartido",
                )
                break
    if not any("model_config.py" in paragraph.text for paragraph in doc.paragraphs):
        for paragraph in doc.paragraphs:
            if "backend_client.py" not in paragraph.text or "Cliente HTTP" not in paragraph.text:
                continue
            anchor = paragraph
            for line in (
                "│       ├── model_config.py           # Configuración ML ligera para clientes",
                "│       ├── file_utils.py             # Escrituras atómicas de secretos y estado",
                "│       ├── network.py                # Política de loopback y bind",
            ):
                anchor = insert_paragraph_after(anchor, line)
            break
    replace_prefix(
        doc,
        "# 3. Ejecutar la aplicación web",
        "# 3. Ejecutar primero el backend central",
    )
    replace_prefix(
        doc,
        "# 3b. Opcional: API HTTP local",
        "# 3b. Ejecutar el cliente web en otra terminal",
    )
    replace_prefix(
        doc,
        "# 3c. Opcional: servidor de la extensión Gmail",
        "# 3c. El proxy 8765 solo se usa por compatibilidad antigua",
    )
    replace_next_nonempty(
        doc,
        "# 3. Ejecutar primero el backend central",
        "$env:PYTHONPATH=\"src\"; python src/backend_server.py",
    )
    replace_next_nonempty(
        doc,
        "# 3b. Ejecutar el cliente web en otra terminal",
        "$env:PYTHONPATH=\"src\"; streamlit run src/app.py",
    )
    replace_next_nonempty(
        doc,
        "# 3c. El proxy 8765 solo se usa por compatibilidad antigua",
        "$env:PYTHONPATH=\"src\"; python src/gmail_extension_server.py",
    )
    annex_heading = "Anexo E: Resumen de arquitectura y evidencia"
    if not any(paragraph.text == annex_heading for paragraph in doc.paragraphs):
        doc.add_paragraph(annex_heading, style="Heading 2")
        doc.add_paragraph(
            "El sistema es cliente-servidor. Streamlit, extensión y monitor consumen por HTTP "
            "el backend central, único proceso que analiza y mantiene las versiones ES/EN. "
            "En la defensa se ejecuta sobre loopback, aunque puede separarse por configuración."
        )
        doc.add_paragraph(
            "La evidencia reproducible comprende 87 pruebas Python, 2 recorridos con Chromium, "
            "integración continua, benchmark, calibración separada de 40 casos, evaluación de "
            "16 EML reservados y diagnóstico de 1.528 textos DIFrauD con riesgo de solapamiento."
        )
    replace_prefix(
        doc,
        "La evidencia reproducible comprende",
        "La evidencia reproducible comprende 87 pruebas Python, 2 recorridos con Chromium, integración continua, benchmark, calibración separada de 40 casos, evaluación de 16 EML reservados y diagnóstico de 1.528 textos DIFrauD. Los EML son sintéticos y DIFrauD conserva riesgo de solapamiento; ninguna cifra estima producción.",
    )
    _keep_table_rows_together(doc)
    if missing:
        raise ValueError(f"No se encontraron párrafos de la memoria: {missing}")
    doc.save(path)


def sync_full_guide() -> None:
    path = ROOT / "Guia_defensa_TFG.docx"
    doc = Document(path)
    replacements = [
        ("Rama de referencia:", "Rama de referencia: main"),
        (
            "Idea clave:",
            "Idea clave: el sistema es cliente-servidor. El navegador usa Streamlit como presentación; Streamlit, la extensión y el monitor son clientes HTTP de backend_server.py. Solo el backend analiza, entrena y mantiene una versión activa por idioma. En local los procesos comparten equipo y loopback, pero no responsabilidades ni memoria.",
        ),
        (
            "Estos resultados proceden del holdout controlado",
            "Estos resultados proceden de 16 EML locales reservados después de calibrar con otros 40 casos. El informe conserva hashes, escenarios y predicciones por caso. Sirve como evidencia funcional, no para estimar producción.",
        ),
        (
            "Streamlit conserva en session_state",
            "Streamlit conserva en session_state la cuenta Gmail y los resultados importados. Si falta un modelo válido, está corrupto o pertenece al idioma incorrecto, el backend crea un fallback sintético del idioma esperado; nunca usa silenciosamente el modelo de la lengua opuesta. La interfaz lo comunica.",
        ),
        (
            "La validación actual",
            "La validación actual ejecuta 87 pruebas Python y 2 recorridos reales con Chromium. Todas pasan. Incluye web-backend, señales BEC, respaldo de defensa y el recorrido local Gmail EML-JSON-HTTP-backend-Telegram. Las pruebas neuronales rápidas pueden generar ConvergenceWarning esperado; no es un fallo funcional.",
        ),
        (
            "La formulación rigurosa es:",
            "La formulación rigurosa es: 'La rama main ejecuta 87 pruebas Python, 2 recorridos Chromium —incluido cliente web a backend—, CI, calibración y evaluaciones reproducibles'. Los 16 EML son sintéticos y DIFrauD puede solaparse con fuentes del modelo; no demuestran eficacia universal.",
        ),
        (
            "El repositorio separa 40 casos de calibración",
            "El repositorio separa 40 casos de calibración y 16 EML finales (4 por idioma y clase). El evaluador genera accuracy, precisión, recall, F1, accuracy balanceada, VP/VN/FP/FN, hashes y detalle por caso. En los EML, el heurístico obtiene 100,0 %; el combinado, 93,8 % de accuracy y 100,0 % de recall; el neuronal, 75,0 % en ambas. DIFrauD añade 1.528 textos: 90,8 % de accuracy y 96,4 % de recall combinados, con riesgo de fuga documentado.",
        ),
        (
            "normalizar_etiqueta",
            "normalizar_etiqueta acepta formas habituales de phishing y correo legítimo, con error para valores desconocidos. La rama main ejecuta 87 pruebas Python, 2 pruebas de navegador, CI, calibración, evaluaciones reproducibles, respaldo de defensa y benchmark.",
        ),
    ]
    replacements.extend(
        [
            (
                "Versión revisada:",
                "Versión revisada: 30 de agosto de 2026",
            ),
            (
                "Aplicación web local modular",
                "Sistema cliente-servidor local con análisis heurístico y clasificador TF-IDF + MLP",
            ),
            (
                "Tesis del trabajo:",
                "Tesis del trabajo: es viable reunir varios clientes de correo con un backend central explicable y entrenable. La heurística aporta trazabilidad, el modelo aprende patrones y todos los clientes comparten una versión activa por idioma. El prototipo demuestra integración y funcionamiento; no sustituye una plataforma empresarial.",
            ),
            (
                "La propuesta del TFG",
                "La propuesta del TFG es un sistema cliente-servidor local. Streamlit organiza la interacción y consume por HTTP un backend obligatorio que centraliza parser, reglas, modelos y entrenamiento. La extensión y el monitor usan el mismo contrato y no duplican ni cargan modelos.",
            ),
            (
                "Frase de apertura:",
                "Frase de apertura: Mi trabajo diseña y evalúa un sistema cliente-servidor para analizar correos mediante reglas explicables y modelos TF-IDF + MLP en español e inglés. Una actualización del servidor se aplica a todos los clientes.",
            ),
            (
                "Usuario\n  -> Streamlit:",
                "Usuario / Gmail\n  -> Clientes: Streamlit | extensión | monitor\n      -> HTTP/JSON: texto | EML Base64 | campos estructurados\n          -> Backend central: parser y normalización\n              -> Rama heurística: señales -> pesos -> explicación\n              -> Rama neuronal: idioma -> TF-IDF -> MLP\n          -> Decisión y versión del modelo\n  <- Respuesta JSON: riesgo, explicación y metadatos\n  -> Presentación o alerta Telegram",
            ),
            (
                "Streamlit reduce el trabajo de frontend",
                "Streamlit reduce el trabajo de frontend y permite dedicar el TFG al motor. Un sistema visual compartido aporta una cabecera de marca, navegación adaptable, estados, tarjetas y formularios coherentes; Detección guía el flujo en tres pasos y Configuración lo organiza en pestañas. En esta arquitectura se limita a presentación y a llamadas con BackendClient; no importa ModelStorage, no ejecuta reglas ni muestra rutas locales completas de credenciales. Cada interacción reejecuta el script de UI y el framework no aporta por sí solo autenticación multiusuario, por lo que sigue siendo apropiado para el prototipo y no basta para publicar un servicio abierto.",
            ),
            (
                "El recorrido siempre intenta",
                "El recorrido siempre atraviesa el mismo contrato HTTP. El cliente conserva únicamente la entrada y el estado visual; el backend produce la representación interna y el resultado completo, evitando que web, extensión y monitor repliquen reglas o decisiones.",
            ),
            (
                "2.\tLa entrada se convierte",
                "2.\tEl cliente envía texto, bytes EML codificados o campos estructurados al endpoint /analyze.",
            ),
            (
                "3.\tCorreoAnalizado.from_input",
                "3.\tEl backend parsea y normaliza remitente, asunto, cuerpo, HTML, cabeceras, URLs, anclas y adjuntos.",
            ),
            (
                "7.\tSi el modo usa el clasificador",
                "7.\tSi el modo usa el clasificador, el backend detecta español o inglés y reutiliza el detector central cacheado por idioma.",
            ),
            (
                "10.\tLa interfaz muestra",
                "10.\tEl backend devuelve riesgo, veredicto, señales, explicación, idioma y versión. El cliente los muestra y el monitor puede alertar por Telegram.",
            ),
            (
                "El pipeline reúne vectorizador y MLP",
                "El pipeline reúne vectorizador y MLP en un solo objeto del servidor. ModelStorage encapsula carga y guardado; backend_service activa cada entrenamiento con sustitución atómica e invalida la caché para que las siguientes peticiones usen la versión nueva.",
            ),
            (
                "La aplicación conserva en el modelo",
                "El backend conserva en el modelo fuentes, columnas, fecha, hiperparámetros y estadísticas, pero no textos crudos. Los clientes solo reciben metadatos y un identificador SHA-256 abreviado de versión; los joblib nunca se transfieren por la API.",
            ),
            (
                "Streamlit conserva en session_state",
                "Streamlit conserva en session_state la cuenta Gmail y resultados para presentarlos. El estado de modelos procede de /health y /models. Si falta un artefacto, está corrupto o pertenece al idioma incorrecto, únicamente el backend crea un fallback sintético del idioma esperado y lo comunica; el cliente no mantiene copias.",
            ),
            (
                "•\tReutilización:",
                "•\tReutilización: detección manual y monitor consumen el mismo contrato y las mismas versiones del backend, sin duplicar reglas.",
            ),
            (
                "No. La aplicación web es la interfaz principal",
                "No. El monitor confirma la arquitectura: obtiene Gmail, envía cada mensaje al backend mediante RemoteAnalysisService y usa la respuesta para decidir la alerta. No carga modelos ni duplica reglas. Gmail y Telegram siguen siendo servicios externos.",
            ),
            (
                "Antes de entrenar, la interfaz resume",
                "Antes de entrenar, el cliente envía los CSV a /datasets/summary. El backend valida columnas y devuelve filas, phishing y legítimos por archivo; el navegador solo presenta el resumen.",
            ),
            (
                "El usuario selecciona idioma",
                "El usuario selecciona idioma, columnas e hiperparámetros. El cliente los envía a /train; el backend crea un clasificador desde cero, registra procedencia, ajusta el pipeline, guarda de forma atómica, invalida la caché y devuelve la versión activa para todos.",
            ),
            (
                "La pestaña Evaluar",
                "La pestaña Evaluar envía un CSV a /evaluate. El backend carga el modelo activo, calcula accuracy, precision, recall, F1, balanced accuracy y matriz de confusión y devuelve las métricas sin modificar la versión.",
            ),
            (
                "La pestaña Comparar",
                "La pestaña Comparar envía entrenamiento, prueba y hasta tres configuraciones a /compare. El backend entrena en memoria y devuelve métricas; ninguna configuración sustituye el modelo activo.",
            ),
            (
                "Reduce complejidad, superficie de ataque",
                "Centraliza decisiones y modelos: los clientes solo transportan entradas y presentan salidas. Cambiar una versión en el backend cambia a todos. El coste es operar un proceso obligatorio y, si se despliega, añadir TLS, autenticación, rate limiting y un registro compartido para réplicas.",
            ),
            (
                "Sí. Streamlit sirve una interfaz web",
                "Sí. Streamlit sirve la interfaz web y actúa como cliente del backend. 'Local' describe que ambos procesos están en el mismo equipo; no elimina la separación cliente-servidor.",
            ),
            (
                "Internamente Streamlit usa HTTP",
                "Sí, Streamlit usa HTTP con el navegador. Además, en este proyecto su proceso Python llama por HTTP al backend central obligatorio. Es un recorrido de dos saltos: navegador -> Streamlit -> backend.",
            ),
            (
                "EmailAnalysisService detecta español",
                "El backend detecta español o inglés por mensaje y mantiene un detector cacheado por idioma. Web, extensión y monitor reciben el idioma y la versión en la respuesta; no seleccionan ni cargan joblib.",
            ),
            (
                "Se crea un modelo sintético",
                "El backend crea un fallback sintético de diez ejemplos para mantener la demo. /health indica que no existe artefacto persistido y los clientes muestran el aviso. No se distribuye ni se confunde con un modelo validado.",
            ),
            (
                "Sí. EmailAnalysisService detecta",
                "Sí. El backend detecta el idioma por mensaje y cachea el modelo correspondiente. El monitor solo usa RemoteAnalysisService, de modo que comparte exactamente las versiones que usa la web.",
            ),
            (
                "Son adaptadores locales opcionales",
                "backend_server.py es el servidor central obligatorio. gmail_extension_server.py queda solo como proxy opcional para instalaciones antiguas del puerto 8765; la extensión actual llama directamente al puerto 8766.",
            ),
            (
                "La documentación, el README y las guías",
                "La documentación, el README y las guías se han alineado con la rama main. El flujo es cliente-servidor: Streamlit, extensión y monitor consumen un backend central que mantiene las versiones ES/EN.",
            ),
            (
                "README.md documenta",
                "README.md documenta el arranque en dos terminales, el contrato HTTP, la extensión directa al puerto 8766, el proxy 8765 opcional y el monitor. Para un destino remoto exige HTTPS; --allow-remote requiere un token administrativo de al menos 24 caracteres y aún se delimitan los controles de producción pendientes.",
            ),
            (
                "La vista Inicio refleja",
                "La vista Inicio consulta /health y muestra el backend y sus versiones. Detección, Entrenamiento y Monitor llaman a BackendClient; solo el servidor selecciona y cachea modelos.",
            ),
            (
                "Entrenamiento y evaluación muestran",
                "Entrenamiento y evaluación presentan las métricas devueltas por el backend. Cada /train comienza con los CSV recibidos, registra procedencia, activa la versión de forma atómica y no serializa textos crudos.",
            ),
            (
                "Existe una API HTTP local opcional",
                "Existe un backend HTTP central obligatorio para web, extensión y monitor; en la configuración académica escucha en loopback.",
            ),
            (
                "Cierre final:",
                "Cierre final: He construido un sistema cliente-servidor que normaliza correos de texto, EML o Gmail, evalúa señales explicables —incluido BEC sin enlace— y aplica modelos TF-IDF + MLP centrales. Streamlit, extensión y monitor comparten las mismas versiones; el backend puede actualizarse sin distribuir modelos. Sus límites son la independencia y actualidad de los datos externos y los controles de producción.",
            ),
            (
                "1.\tEjecuta streamlit run",
                "1.\tEjecuta python src/backend_server.py, comprueba /health y abre Streamlit en otra terminal.",
            ),
            (
                "8.\tCierra recordando",
                "8.\tCierra recordando que Gmail y Telegram son opcionales, pero el backend central es obligatorio y compartido.",
            ),
            (
                "•\tSi Streamlit tarda",
                "•\tSi backend o Streamlit fallan, muestra /health, una respuesta JSON o el recorrido E2E automatizado guardado.",
            ),
            (
                "TLS, autenticación de inferencia y administración",
                "TLS, autenticación de inferencia y administración, gestión de secretos, rate limiting, logging seguro, almacenamiento de modelos compartido para réplicas y confirmación externa con datos reales recientes.",
            ),
        ]
    )
    missing = []
    for prefix, replacement in replacements:
        if not replace_prefix(doc, prefix, replacement) and not any(
            paragraph.text == replacement for paragraph in _all_paragraphs(doc)
        ):
            missing.append(prefix)
    replace_fragment(doc, "47 pruebas", "87 pruebas Python y 2 de navegador")
    replace_fragment(doc, "54 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "59 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "60 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "72 pruebas Python", "87 pruebas Python")
    replace_fragment(doc, "81 pruebas Python", "87 pruebas Python")
    replace_fragment(
        doc,
        "¿Por qué pesos 60/40?",
        "¿Por qué pesos 35/65 y alta confianza 70?",
    )
    replace_fragment(
        doc,
        "por defecto 60 % heurística y 40 % red neuronal",
        "por defecto 35 % heurística y 65 % neuronal, con alta confianza 70",
    )
    replace_fragment(
        doc,
        "configuración habitual: 60 % heurística + 40 % neuronal",
        "configuración calibrada: 35 % heurística + 65 % neuronal; alta confianza: 70",
    )
    replace_fragment(
        doc,
        "¿Por qué pesos 20/80 y alta confianza 70?",
        "¿Por qué pesos 35/65 y alta confianza 70?",
    )
    replace_fragment(
        doc,
        "por defecto 20 % heurística y 80 % neuronal, con alta confianza 70",
        "por defecto 35 % heurística y 65 % neuronal, con alta confianza 70",
    )
    replace_fragment(
        doc,
        "configuración calibrada: 20 % heurística + 80 % neuronal; alta confianza: 70",
        "configuración calibrada: 35 % heurística + 65 % neuronal; alta confianza: 70",
    )
    replace_fragment(
        doc,
        "¿Por qué pesos 30/70 y alta confianza 70?",
        "¿Por qué pesos 35/65 y alta confianza 70?",
    )
    replace_fragment(
        doc,
        "por defecto 30 % heurística y 70 % neuronal, con alta confianza 70",
        "por defecto 35 % heurística y 65 % neuronal, con alta confianza 70",
    )
    replace_fragment(
        doc,
        "configuración calibrada: 30 % heurística + 70 % neuronal; alta confianza: 70",
        "configuración calibrada: 35 % heurística + 65 % neuronal; alta confianza: 70",
    )
    replace_fragment(
        doc,
        "El umbral 45 es una decisión operativa del prototipo. No procede de una calibración estadística exhaustiva y debe presentarse como parámetro ajustable. Para producción habría que optimizarlo con una curva precision-recall y con costes explícitos de falsos positivos y falsos negativos.",
        "El umbral 26 fue seleccionado por la rejilla estratificada de 40 casos junto a la fusión 35/65. Sigue siendo un parámetro ajustable: para producción habría que confirmarlo sobre datos reales independientes, con una curva precision-recall y costes explícitos de falsos positivos y falsos negativos.",
    )
    replace_fragment(
        doc,
        "El umbral 45 fue seleccionado por la rejilla estratificada de 40 casos junto a la fusión 20/80. Sigue siendo un parámetro ajustable: para producción habría que confirmarlo sobre datos reales representativos, con una curva precision-recall y costes explícitos de falsos positivos y falsos negativos.",
        "El umbral 26 fue seleccionado por la rejilla estratificada de 40 casos junto a la fusión 35/65. Sigue siendo un parámetro ajustable: para producción habría que confirmarlo sobre datos reales independientes, con una curva precision-recall y costes explícitos de falsos positivos y falsos negativos.",
    )
    replace_fragment(
        doc,
        "El umbral 36 fue seleccionado por la rejilla estratificada de 40 casos junto a la fusión 30/70. Sigue siendo un parámetro ajustable: para producción habría que confirmarlo sobre datos reales independientes, con una curva precision-recall y costes explícitos de falsos positivos y falsos negativos.",
        "El umbral 26 fue seleccionado por la rejilla estratificada de 40 casos junto a la fusión 35/65. Sigue siendo un parámetro ajustable: para producción habría que confirmarlo sobre datos reales independientes, con una curva precision-recall y costes explícitos de falsos positivos y falsos negativos.",
    )
    replace_fragment(
        doc,
        "El ajuste sobre entrenamiento es casi perfecto, pero el holdout controlado no es representativo y no permite generalizar a producción.",
        "El ajuste sobre entrenamiento es casi perfecto, pero los EML sintéticos reservados no son una muestra representativa y no permiten generalizar a producción.",
    )
    replace_prefix(
        doc,
        "Es un umbral operativo del prototipo",
        "El valor 26 fue seleccionado por la rejilla estratificada junto a la fusión 35/65. Debe confirmarse con un corpus real independiente y con el coste concreto de falsos positivos y negativos.",
    )
    replace_prefix(
        doc,
        "El valor 45 se mantuvo porque la rejilla estratificada",
        "El valor 26 fue seleccionado por la rejilla estratificada junto a la fusión 35/65. Debe confirmarse con un corpus real independiente y con el coste concreto de falsos positivos y negativos.",
    )
    replace_prefix(
        doc,
        "El valor 36 fue seleccionado por la rejilla estratificada",
        "El valor 26 fue seleccionado por la rejilla estratificada junto a la fusión 35/65. Debe confirmarse con un corpus real independiente y con el coste concreto de falsos positivos y negativos.",
    )
    replace_prefix(
        doc,
        "Se prioriza ligeramente la trazabilidad heurística",
        "La rejilla de calibración sobre 40 casos seleccionó 35 % heurístico y 65 % neuronal, reservando al menos 20 % a cada detector. El nivel 70 conserva una evidencia individual concluyente. Los 16 EML posteriores se evaluaron sin reajustar esos valores.",
    )
    replace_prefix(
        doc,
        "La rejilla de calibración sobre 40 casos seleccionó 20 % heurístico",
        "La rejilla de calibración sobre 40 casos seleccionó 35 % heurístico y 65 % neuronal, reservando al menos 20 % a cada detector. El nivel 70 conserva una evidencia individual concluyente. Los 16 EML posteriores se evaluaron sin reajustar esos valores.",
    )
    replace_prefix(
        doc,
        "La rejilla de calibración sobre 40 casos seleccionó 30 % heurístico",
        "La rejilla de calibración sobre 40 casos seleccionó 35 % heurístico y 65 % neuronal, reservando al menos 20 % a cada detector. El nivel 70 conserva una evidencia individual concluyente. Los 16 EML posteriores se evaluaron sin reajustar esos valores.",
    )
    replace_fragment(doc, "26 de agosto de 2026", "30 de agosto de 2026")
    replace_fragment(doc, "27 de agosto de 2026", "30 de agosto de 2026")
    replace_fragment(doc, "28 de agosto de 2026", "30 de agosto de 2026")
    replace_fragment(doc, "29 de agosto de 2026", "30 de agosto de 2026")
    replace_fragment(doc, "propuesta híbrida", "arquitectura cliente-servidor")
    replace_fragment(
        doc,
        "Flujo local modular; Streamlit y adaptadores opcionales.",
        "Clientes HTTP y backend central con modelos compartidos.",
    )
    replace_fragment(
        doc,
        "src/app.py, *_app.py, ui_components.py",
        "src/app.py, *_app.py, backend_client.py",
    )
    replace_fragment(
        doc,
        "Navegación, formularios, estado visual y presentación del resultado.",
        "Navegación, formularios, llamadas HTTP y presentación.",
    )
    replace_fragment(
        doc,
        "detect_app.py, gmail_monitor.py, analysis_service.py",
        "backend_service.py, analysis_service.py, http_api.py",
    )
    replace_fragment(
        doc,
        "Coordinar análisis, selección de modo, combinación y monitorización.",
        "Coordinar análisis, modos, entrenamiento, versiones y transporte.",
    )
    replace_fragment(
        doc,
        ".env.local, token.json, estado_monitor.json, *.joblib",
        "Cliente: .env.local/token/estado; servidor: *.joblib",
    )
    replace_fragment(
        doc,
        "El monitor detecta el idioma por mensaje mediante el servicio común y aplica un fallback controlado.",
        "El backend detecta el idioma por mensaje; el monitor consume esa decisión remota y el fallback controlado.",
    )
    replace_fragment(doc, "feature/web", "main")
    replace_fragment(doc, "28 señales", "31 señales")
    replace_fragment(doc, "28 reglas", "31 reglas")
    replace_fragment(doc, "superar 45", "superar 26")
    replace_fragment(doc, "superar 36", "superar 26")
    replace_fragment(doc, "phishing_heuristico = riesgo_heuristico >= 45", "phishing_heuristico = riesgo_heuristico >= 26")
    replace_fragment(doc, "phishing_heurístico = riesgo_heurístico >= 45", "phishing_heurístico = riesgo_heurístico >= 26")
    replace_fragment(
        doc,
        "Los 27 pesos positivos suman 1,97; por tanto, varios indicadores pueden saturar la puntuación y el límite evita superar 100.",
        "Los 30 pesos positivos suman 2,21. Además, la coincidencia de los tres indicios BEC añade un refuerzo de 0,46; los indicios aislados conservan peso bajo. El límite evita superar 100.",
    )
    replace_fragment(
        doc,
        "Urgencia, saludo genérico, credenciales, asunto, referencias a archivos.",
        "Urgencia, saludo genérico, credenciales, asunto, referencias a archivos y patrones BEC.",
    )
    replace_fragment(doc, "umbral habitual del combinado: 45", "umbral calibrado del combinado: 26")
    replace_fragment(doc, "¿Por qué un umbral de 45?", "¿Por qué un umbral de 26?")
    replace_fragment(doc, "Combinado con 60 % heurística", "Combinado con 35 % heurística")
    replace_fragment(
        doc,
        "Si no hay modelo del idioma esperado, se prueba el del otro idioma",
        "Si no hay modelo del idioma esperado, se crea un fallback sintético de ese mismo idioma",
    )
    replace_prefix(
        doc,
        "Interpretación correcta:",
        "Interpretación correcta: 40 casos controlados calibran y 16 EML distintos evalúan. "
        "El combinado obtiene 93,8 % de accuracy y 100,0 % de recall en esos 16 escenarios; DIFrauD añade 90,8 % de accuracy y 96,4 % de recall, pero "
        "la primera muestra es sintética y la segunda puede solaparse con fuentes del entrenamiento.",
    )
    replace_fragment(
        doc,
        "la evaluación independiente es pequeña y no permite generalizar",
        "los corpus sintéticos de calibración y EML no son representativos y no permiten generalizar a producción",
    )
    replace_fragment(
        doc,
        "pip install -r requirements.txt",
        "python -m pip install -r requirements-dev.txt -c constraints.txt",
    )
    for paragraph in _all_paragraphs(doc):
        if not paragraph.text.startswith("python -m venv .venv"):
            continue
        commands = paragraph.text
        if "python src/backend_server.py" not in commands:
            commands = commands.replace(
                "streamlit run src/app.py",
                '$env:PYTHONPATH="src"; python src/backend_server.py\n'
                '$env:PYTHONPATH="src"; streamlit run src/app.py',
            )
        _set_text(paragraph, commands)
        break
    replace_prefix(
        doc,
        "•\tConozco las inconsistencias actuales del repositorio",
        "•\tConozco las limitaciones actuales del prototipo y no hago afirmaciones que excedan la evidencia.",
    )
    replace_prefix(
        doc,
        "•\tHe decidido si corregir o retirar los artefactos heredados",
        "•\tHe comprobado que la documentación, las guías y la evidencia reproducible están sincronizadas.",
    )
    _restore_page_fields(doc)
    _replace_controlled_results_table(doc)
    _replace_signal_weights_table(doc)
    _keep_table_rows_together(doc)
    if missing:
        raise ValueError(f"No se encontraron párrafos de la guía: {missing}")
    doc.save(path)


def iter_blocks(doc: DocumentObject) -> Iterable[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def export_text(docx_path: Path, text_path: Path) -> None:
    doc = Document(docx_path)
    lines = []
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            lines.append(block.text.rstrip())
        else:
            for row in block.rows:
                lines.append(
                    "\t".join(cell.text.replace("\n", " ") for cell in row.cells).rstrip()
                )
        lines.append("")
    text_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def main() -> None:
    sync_memory()
    sync_full_guide()
    export_text(ROOT / "TFG.docx", ROOT / "TFG.txt")
    export_text(ROOT / "Guia_defensa_TFG.docx", ROOT / "Guia_defensa_TFG.txt")
    print("Memoria y guía extensa sincronizadas.")


if __name__ == "__main__":
    main()
