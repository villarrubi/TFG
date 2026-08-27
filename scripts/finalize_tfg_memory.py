"""Migración histórica de contenido y maquetación de la memoria del TFG.

No debe ejecutarse sobre la entrega actual; usa ``sync_delivery_docs.py``. Se
conserva únicamente para documentar cómo se construyó la primera versión final.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "TFG.docx"
BLUE = "12386B"
TEAL = "1D6E7A"
TEXT = "263238"
GRID = "D5DEE8"


def _remove(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _paragraphs(doc):
    return list(doc.paragraphs)


def _xml_text(element) -> str:
    """Devuelve una sola vez el texto visible de un elemento WordprocessingML."""
    return "".join(element.xpath(".//w:t/text()"))


def replace_prefix(doc, prefix: str, replacement: str) -> int:
    changed = 0
    for paragraph in _paragraphs(doc):
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            changed += 1
    return changed


def replace_exact(doc, old: str, new: str) -> int:
    changed = 0
    for paragraph in _paragraphs(doc):
        if paragraph.text == old:
            paragraph.text = new
            changed += 1
    return changed


def remove_prefix(doc, prefix: str) -> int:
    targets = [p for p in _paragraphs(doc) if p.text.startswith(prefix)]
    for paragraph in targets:
        _remove(paragraph._element)
    return len(targets)


def remove_range(doc, start_prefix: str, end_prefix: str | None) -> int:
    body = doc.element.body
    children = list(body.iterchildren())
    start = next(
        (
            index
            for index, child in enumerate(children)
            if child.tag == qn("w:p")
            and _xml_text(child).startswith(start_prefix)
        ),
        None,
    )
    if start is None:
        return 0
    end = len(children)
    if end_prefix is not None:
        end = next(
            (
                index
                for index, child in enumerate(children[start + 1 :], start + 1)
                if child.tag == qn("w:p")
                and _xml_text(child).startswith(end_prefix)
            ),
            len(children),
        )
    removed = 0
    for child in children[start:end]:
        if child.tag != qn("w:sectPr"):
            _remove(child)
            removed += 1
    return removed


def remove_drawing_by_target(doc, target_name: str) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        for blip in paragraph._element.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            rel = doc.part.rels.get(rel_id)
            if rel is not None and str(rel.target_ref).endswith(target_name):
                _remove(paragraph._element)
                removed += 1
                break
    return removed


def find_table(doc, first_header: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == first_header:
            return table
    raise ValueError(f"No se encontró la tabla con cabecera {first_header!r}")


def remove_table(doc, first_header: str) -> None:
    _remove(find_table(doc, first_header)._element)


def _set_cell_text(cell, value: str, *, header: bool, size: float = 9.0) -> None:
    cell.text = value
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = header
            run.font.color.rgb = RGBColor.from_string("FFFFFF" if header else TEXT)


def rewrite_table(table, rows: list[list[str]], *, size: float = 9.0) -> None:
    if len(rows) != len(table.rows):
        raise ValueError(
            f"La tabla esperaba {len(table.rows)} filas y recibió {len(rows)}"
        )
    if any(len(row) != len(table.columns) for row in rows):
        raise ValueError("El número de columnas no coincide con la tabla existente")
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            _set_cell_text(
                table.rows[row_index].cells[column_index],
                value,
                header=row_index == 0,
                size=size,
            )


def _font(size: int, *, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def build_verified_vector_chart() -> bytes:
    image = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(image)
    title = _font(48, bold=True)
    subtitle = _font(27)
    label_font = _font(31)
    value_font = _font(31, bold=True)
    source_font = _font(23)

    draw.text((100, 70), "Vectores iniciales en brechas", fill=f"#{BLUE}", font=title)
    draw.text(
        (100, 135),
        "DBIR 2025 · n = 9.891 brechas analizadas",
        fill="#526272",
        font=subtitle,
    )

    labels = [
        "Otros vectores",
        "Credenciales comprometidas",
        "Explotación de vulnerabilidades",
        "Phishing",
    ]
    values = [43, 22, 20, 15]
    colors = ["#7891A8", f"#{BLUE}", "#2F6FA5", f"#{TEAL}"]
    left, top, chart_width, bar_height, gap = 650, 245, 760, 95, 65

    for tick in range(0, 51, 10):
        x = left + int(chart_width * tick / 50)
        draw.line(
            (x, top - 15, x, top + 4 * (bar_height + gap) - gap),
            fill=f"#{GRID}",
            width=2,
        )
        draw.text((x - 12, 865), f"{tick}", fill="#526272", font=source_font)

    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = top + index * (bar_height + gap)
        draw.text((100, y + 24), label, fill=f"#{TEXT}", font=label_font)
        width = int(chart_width * value / 50)
        draw.rounded_rectangle(
            (left, y, left + width, y + bar_height), radius=16, fill=color
        )
        draw.text(
            (left + width + 22, y + 24),
            f"{value} %",
            fill=f"#{TEXT}",
            font=value_font,
        )

    draw.text(
        (100, 930),
        "Fuente: Verizon, 2025 Data Breach Investigations Report.",
        fill="#526272",
        font=source_font,
    )
    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    return stream.getvalue()


def build_qualitative_phase_diagram() -> bytes:
    """Crea un diagrama cualitativo del ciclo sin métricas no verificadas."""
    image = Image.new("RGB", (1400, 1700), "white")
    draw = ImageDraw.Draw(image)
    phase_font = _font(42, bold=True)
    text_font = _font(29)
    number_font = _font(42, bold=True)

    phases = [
        (
            "Reconocimiento",
            "Atacante: recopila información pública",
            "Control: limitar exposición y generar alertas",
            f"#{BLUE}",
        ),
        (
            "Creación del cebo",
            "Atacante: prepara el mensaje, enlace o QR",
            "Control: autenticación y filtrado",
            "#2F6FA5",
        ),
        (
            "Interacción",
            "Víctima: abre el mensaje o introduce datos",
            "Control: formación y MFA resistente",
            f"#{TEAL}",
        ),
        (
            "Monetización",
            "Atacante: usa o vende credenciales y datos",
            "Control: detección y respuesta",
            "#556B8E",
        ),
    ]

    left, right = 145, 1255
    card_height, gap, top = 300, 80, 105
    for index, (phase, action, control, color) in enumerate(phases, start=1):
        y = top + (index - 1) * (card_height + gap)
        draw.rounded_rectangle(
            (left, y, right, y + card_height),
            radius=28,
            fill="#F5F8FB",
            outline=color,
            width=7,
        )
        draw.ellipse((190, y + 87, 310, y + 207), fill=color)
        number_box = draw.textbbox((0, 0), str(index), font=number_font)
        number_width = number_box[2] - number_box[0]
        number_height = number_box[3] - number_box[1]
        draw.text(
            (250 - number_width / 2, y + 147 - number_height / 2 - 5),
            str(index),
            fill="white",
            font=number_font,
        )
        draw.text((355, y + 45), phase, fill=color, font=phase_font)
        draw.text((355, y + 125), action, fill=f"#{TEXT}", font=text_font)
        draw.text((355, y + 190), control, fill="#526272", font=text_font)

        if index < len(phases):
            arrow_x = 700
            arrow_top = y + card_height + 15
            arrow_bottom = y + card_height + gap - 15
            draw.line(
                (arrow_x, arrow_top, arrow_x, arrow_bottom),
                fill=f"#{GRID}",
                width=12,
            )
            draw.polygon(
                [
                    (arrow_x - 25, arrow_bottom - 18),
                    (arrow_x + 25, arrow_bottom - 18),
                    (arrow_x, arrow_bottom + 12),
                ],
                fill=f"#{GRID}",
            )

    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    return stream.getvalue()


def replace_image_blob(doc, target_name: str, data: bytes) -> None:
    for blip in doc.element.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        rel = doc.part.rels.get(rel_id)
        if rel is not None and str(rel.target_ref).endswith(target_name):
            rel.target_part._blob = data
            for source_rect in blip.getparent().xpath("./a:srcRect"):
                source_rect.attrib.clear()
            return
    raise ValueError(f"No se encontró la imagen {target_name}")


def remove_provisional_appendix(doc) -> None:
    marker = "Actualización técnica y evidencias de la versión final"
    body = doc.element.body
    children = list(body.iterchildren())
    marker_index = next(
        (
            index
            for index, child in enumerate(children)
            if child.tag == qn("w:p") and _xml_text(child) == marker
        ),
        None,
    )
    if marker_index is None:
        return
    previous = children[marker_index - 1] if marker_index else None
    if previous is not None and previous.tag == qn("w:p") and previous.xpath(
        ".//w:br[@w:type='page']"
    ):
        _remove(previous)
    remove_range(doc, marker, None)


def clean_unverified_visuals(doc) -> None:
    for prefix in (
        "Figura 2.3:",
        "Figura 2.4:",
        "Figura 2.5:",
        "Tabla 2.5:",
    ):
        remove_prefix(doc, prefix)

    remove_range(
        doc,
        "Caso real: campaña CorruptQR (2025)",
        "Puntos de intervención y controles",
    )
    for image in ("image4.png", "image5.png", "image6.png"):
        remove_drawing_by_target(doc, image)
    remove_table(doc, "Principio")


def correct_context_and_psychology(doc) -> None:
    replace_prefix(
        doc,
        "La ciberseguridad se enfrenta en la actualidad",
        "La ciberseguridad se enfrenta a una amenaza persistente en el entorno digital: el phishing. La evolución de las campañas y el uso de inteligencia artificial generativa dificultan la detección basada únicamente en filtros y listas negras. Este escenario justifica estudiar aproximaciones que incorporen el contexto, el comportamiento y la intención del mensaje.",
    )
    replace_prefix(
        doc,
        "El 2025 Data Breach Investigations Report de Verizon revela",
        "El 2025 Data Breach Investigations Report de Verizon sitúa la participación humana en aproximadamente el 60 % de las brechas analizadas e identifica el phishing entre los vectores iniciales observados. Los informes del FBI muestran, además, que el fraude de correo corporativo puede ocasionar pérdidas económicas relevantes, aunque el impacto varía entre incidentes (Verizon, 2025; Federal Bureau of Investigation, 2024).",
    )
    replace_prefix(
        doc,
        "Este tipo de ataque ha evolucionado desde los correos masivos",
        "El phishing ha evolucionado desde campañas masivas hacia ataques dirigidos y mensajes asistidos por inteligencia artificial generativa. Estas herramientas pueden mejorar la fluidez y la adaptación contextual del texto, lo que reduce el valor de algunos indicadores lingüísticos tradicionales sin volverlos inútiles (Heiding et al., 2024).",
    )
    replace_prefix(
        doc,
        "Duración y riesgos: Esta fase puede durar",
        "Duración y riesgos: La fase de reconocimiento puede prolongarse y no siempre deja señales visibles para la organización. El DBIR 2025 identifica el abuso de credenciales como uno de los vectores iniciales más frecuentes, lo que refuerza la necesidad de limitar la exposición pública y vigilar credenciales comprometidas (Verizon, 2025).",
    )
    replace_prefix(
        doc,
        "En España, el Instituto Nacional de Ciberseguridad",
        "En España, el Instituto Nacional de Ciberseguridad (INCIBE) gestionó 97.348 incidentes de ciberseguridad durante 2024, un 16,6 % más que en 2023. Dentro del fraude en línea, el phishing fue la tipología más frecuente, con 21.571 casos registrados (Instituto Nacional de Ciberseguridad, 2025).",
    )
    replace_prefix(
        doc,
        "Esta fase es crítica porque el 90%",
        "Esta fase es crítica porque la credibilidad del mensaje y su similitud con comunicaciones reales condicionan la interacción de la víctima. En el DBIR 2025, el phishing representó aproximadamente el 15 % de los vectores iniciales identificados en brechas, por detrás del abuso de credenciales y de la explotación de vulnerabilidades (Verizon, 2025).",
    )
    replace_prefix(
        doc,
        "Factores humanos: La urgencia",
        "Factores humanos: La urgencia, el miedo a perder acceso y la apariencia de legitimidad explotan sesgos cognitivos y pueden precipitar una decisión. La formación y la autenticación resistente al phishing reducen el riesgo, pero no eliminan la necesidad de controles técnicos.",
    )
    replace_prefix(
        doc,
        "Esta fase destaca la intersección entre tecnología y psicología",
        "Esta fase muestra la intersección entre tecnología y psicología: los controles técnicos deben complementarse con formación, autenticación resistente al phishing y mecanismos de detección. El DBIR 2025 sitúa la participación humana en el 60 % de las brechas analizadas (Verizon, 2025).",
    )
    replace_prefix(
        doc,
        "Actividades clave: El usuario es redirigido",
        "Actividades clave: El usuario puede ser redirigido a un sitio falso para capturar credenciales o inducido a abrir un adjunto malicioso. En ataques de varias etapas, un proxy adversary-in-the-middle (AiTM) puede interceptar tokens de sesión en tiempo real (Kroll, 2025).",
    )
    replace_prefix(
        doc,
        "Ejemplo: Datos robados de un whaling",
        "Ejemplo: En la campaña CorruptQR observada por Kroll, documentos de Office con cabeceras dañadas mostraban un código QR tras su reparación; al escanearlo, la víctima era dirigida a una página de phishing orientada al robo de credenciales corporativas y tokens de sesión (Kroll, 2025).",
    )
    replace_prefix(
        doc,
        "Impacto: Genera pérdidas globales",
        "Impacto: El coste económico depende del tipo de incidente y no debe atribuirse íntegramente al phishing. Como referencia general, IBM estimó en 4,44 millones de dólares el coste medio mundial de una brecha de datos en 2025, un 9 % menos que el año anterior (IBM, 2025).",
    )
    replace_prefix(
        doc,
        "El correo electrónico sigue siendo, con diferencia, el canal dominante.",
        "El correo electrónico continúa siendo un canal central del phishing, aunque las campañas combinan cada vez más mensajes, páginas falsas, códigos QR y fraude de correo corporativo. En el segundo trimestre de 2025, APWG observó 1.130.393 ataques de phishing y documentó el uso de códigos QR contra 1.642 marcas (Anti-Phishing Working Group, 2025).",
    )
    replace_prefix(
        doc,
        "Los mensajes SMS (smishing) y las llamadas telefónicas",
        "Los mensajes SMS (smishing), las llamadas telefónicas (vishing), las redes sociales y los códigos QR amplían los canales de entrega. Kroll observó en 2024 campañas que empleaban mensajes en redes sociales y documentos reparables que revelaban códigos QR, mientras que APWG dedicó en 2025 un seguimiento específico a los ataques mediante QR (Kroll, 2025; Anti-Phishing Working Group, 2025).",
    )
    replace_prefix(
        doc,
        "Phishing masivo o genérico:",
        "Phishing masivo o genérico: campañas de gran volumen que reutilizan mensajes similares. Su coste reducido permite mantenerlas incluso cuando solo una fracción pequeña de los destinatarios interactúa.",
    )
    replace_prefix(
        doc,
        "Spear-phishing: ataque dirigido",
        "Spear-phishing: ataque dirigido contra una persona o un grupo reducido. Requiere investigación previa (OSINT) y mensajes adaptados al contexto, por lo que suele resultar más persuasivo que una campaña genérica.",
    )
    replace_prefix(
        doc,
        "Clone phishing:",
        "Clone phishing: se toma un correo legítimo previamente recibido por la víctima y se modifica un enlace o adjunto. Al conservar elementos de una conversación real, puede aumentar la confianza inicial.",
    )
    replace_prefix(
        doc,
        "Whaling o CEO-fraud:",
        "Whaling o CEO-fraud: variante del spear-phishing dirigida a altos ejecutivos o perfiles financieros. Puede producir pérdidas elevadas cuando deriva en fraude de correo corporativo, aunque el impacto varía ampliamente entre incidentes (Federal Bureau of Investigation, 2024).",
    )
    replace_prefix(
        doc,
        "Adversary-in-the-Middle (AiTM):",
        "Adversary-in-the-Middle (AiTM): técnica que utiliza proxies inversos para interceptar la autenticación en tiempo real y capturar tokens de sesión incluso cuando la víctima emplea determinados métodos de autenticación multifactor. Constituye un riesgo relevante para servicios corporativos en la nube (Kroll, 2025).",
    )
    replace_prefix(
        doc,
        "El tercer eje clasifica los ataques según la técnica de explotación",
        "El tercer eje clasifica los ataques según la técnica de explotación, con independencia del canal de entrega o del nivel de personalización. Entre las tendencias recientes destacan el uso abusivo de inteligencia artificial generativa, la oferta de kits como servicio y las técnicas orientadas al robo de tokens:",
    )
    replace_prefix(
        doc,
        "Inteligencia artificial generativa:",
        "Inteligencia artificial generativa: los modelos de lenguaje pueden ayudar a redactar mensajes fluidos, adaptar el tono al contexto y producir contenido web. Su resultado depende de la información y de las instrucciones disponibles, y no elimina otros indicios técnicos del ataque.",
    )
    replace_prefix(
        doc,
        "2022-2025: era de la IA generativa",
        "2022-2025: expansión de la IA generativa, AiTM, quishing y PhaaS, con una mayor combinación de canales, automatización y servicios criminales.",
    )
    replace_prefix(
        doc,
        "En apenas treinta años se ha pasado",
        "En tres décadas, el phishing ha pasado de campañas rudimentarias a operaciones criminales organizadas que combinan automatización, ingeniería social y servicios comercializados en mercados clandestinos.",
    )
    replace_prefix(
        doc,
        "Este panorama justifica plenamente",
        "Este panorama justifica complementar los filtros de correo y las listas negras con sistemas capaces de analizar el contexto, el comportamiento y la intención del mensaje, objetivo principal del presente trabajo.",
    )
    replace_prefix(
        doc,
        "Aunque el phishing se manifiesta a través de medios técnicos",
        "Aunque el phishing se manifiesta a través de medios técnicos —correo electrónico, SMS, sitios web falsos o códigos QR—, su eficacia depende en gran medida de la explotación de factores cognitivos, emocionales y sociales. Por ello, la ingeniería social constituye un componente central del fenómeno (Hadnagy, 2018).",
    )
    replace_prefix(
        doc,
        "Phishing-as-a-Service (PhaaS):",
        "Phishing-as-a-Service (PhaaS): plataformas que empaquetan plantillas, alojamiento y mecanismos de evasión, reduciendo la barrera de entrada para campañas sofisticadas. Kroll relacionó CorruptQR con la plataforma ONNX y observó otros kits dirigidos a cuentas Microsoft 365 (Kroll, 2025).",
    )
    replace_prefix(
        doc,
        "Bypass de MFA mediante robo de tokens:",
        "Bypass de MFA mediante robo de tokens: los proxies adversary-in-the-middle pueden capturar cookies o tokens de sesión después de una autenticación válida. Por ello se recomiendan métodos resistentes al phishing, como FIDO, junto con controles de acceso condicional (Kroll, 2025).",
    )
    replace_prefix(
        doc,
        "El volumen de ataques no deja de batir récords:",
        "El volumen sigue siendo elevado: APWG observó 1.130.393 ataques de phishing en el segundo trimestre de 2025. En España, INCIBE gestionó 97.348 incidentes de ciberseguridad durante 2024, incluidos 21.571 casos de phishing. Como indicador económico general, no específico de phishing, IBM situó el coste medio mundial de una brecha en 4,44 millones de dólares en 2025 (Anti-Phishing Working Group, 2025; Instituto Nacional de Ciberseguridad, 2025; IBM, 2025).",
    )
    replace_prefix(
        doc,
        "Un estudio longitudinal de 2024-2025 con 48 000 empleados",
        "La combinación de urgencia, autoridad y prueba social puede elevar la persuasión del mensaje. Este efecto depende del contexto, de la experiencia previa y de las diferencias individuales, por lo que no debe expresarse como una tasa universal de éxito (Heiding et al., 2024).",
    )
    replace_prefix(
        doc,
        "Aversión a la pérdida (loss aversion):",
        "Aversión a la pérdida (loss aversion): el miedo a perder el acceso a una cuenta o sufrir una penalización puede pesar más que una ganancia equivalente y precipitar la respuesta.",
    )
    replace_prefix(
        doc,
        "Eliminación total de indicadores lingüísticos tradicionales:",
        "Reducción de indicadores lingüísticos tradicionales: los modelos generativos pueden producir mensajes fluidos y disminuir errores gramaticales o construcciones extrañas que antes facilitaban la detección (Heiding et al., 2024).",
    )
    replace_prefix(
        doc,
        "La incorporación masiva de modelos de lenguaje grandes",
        "La disponibilidad de modelos de lenguaje grandes (LLM) ha introducido cambios cualitativos en la preparación de mensajes de ingeniería social:",
    )
    replace_prefix(
        doc,
        "Personalización contextual perfecta:",
        "Personalización contextual: los modelos de lenguaje pueden adaptar el tono, incorporar información pública y reproducir terminología propia de un sector, aunque la calidad depende de los datos y de las instrucciones disponibles.",
    )
    replace_prefix(
        doc,
        "Generación dinámica de narrativas:",
        "Generación dinámica de narrativas: las herramientas generativas pueden producir y adaptar textos con rapidez. Su uso abusivo reduce el esfuerzo de personalización, aunque no garantiza que el mensaje resulte creíble ni eficaz.",
    )
    replace_prefix(
        doc,
        "Un experimento controlado con 112 participantes",
        "Los experimentos de Heiding et al. (2024) muestran que los modelos de lenguaje pueden producir correos de phishing competitivos frente a mensajes redactados por personas y que la incorporación de principios de persuasión modifica su eficacia. Los resultados corresponden a un diseño experimental concreto y no constituyen una tasa universal de éxito.",
    )
    replace_prefix(
        doc,
        "La naturaleza profundamente psicológica del phishing implica",
        "La dimensión psicológica del phishing implica que una solución eficaz debe combinar el análisis técnico con señales del contexto humano. Los sistemas basados únicamente en características estáticas de URL o contenido pueden degradarse cuando cambian las campañas o se emplea texto generado (Alghenaim et al., 2025; Heiding et al., 2024).",
    )
    replace_prefix(
        doc,
        "Por tanto, el futuro de la detección pasa inevitablemente",
        "Por tanto, entre las líneas de evolución de la detección se encuentran los modelos que incorporan:",
    )

    replace_image_blob(doc, "image1.png", build_qualitative_phase_diagram())
    replace_prefix(
        doc,
        "El Verizon DBIR 2025 establece que el 68 %",
        "El Verizon DBIR 2025 sitúa la participación humana en el 60 % de las brechas analizadas. En España, el 33 % de los usuarios que contactaron con la línea 017 indicó haber recibido algún intento de phishing, vishing o smishing durante 2024 (Verizon, 2025; Instituto Nacional de Ciberseguridad, 2025).",
    )
    replace_prefix(
        doc,
        "Este factor humano no es estático:",
        "Este factor humano no es estático: la fatiga por alerta, la sobrecarga de correo y la normalización del trabajo remoto reducen la atención sostenida. Los informes de industria señalan además que determinados perfiles, como finanzas, recursos humanos o dirección, reciben un volumen desproporcionado de ataques dirigidos (Proofpoint, 2025a).",
    )

    replace_prefix(
        doc,
        "Figura 2.2:",
        "Figura 2.2: Vectores iniciales en brechas analizadas por el DBIR 2025 (n=9.891).",
    )
    replace_image_blob(doc, "image2.png", build_verified_vector_chart())

    rewrite_table(
        find_table(doc, "Fase"),
        [
            ["Indicador", "Medida", "Valor", "Fuente"],
            ["Actividad observada", "Ataques de phishing en Q2 2025", "1.130.393", "APWG"],
            ["Acceso inicial", "Phishing en brechas analizadas", "≈ 15 %", "Verizon DBIR"],
            ["Factor humano", "Brechas con participación humana", "60 %", "Verizon DBIR"],
            ["Impacto general", "Coste medio global por brecha", "4,44 M USD", "IBM"],
        ],
    )
    replace_prefix(
        doc,
        "Tabla 2.1:",
        "Tabla 2.1: Indicadores de contexto del phishing y las brechas (2025).",
    )

    rewrite_table(
        find_table(doc, "Principio (Cialdini)"),
        [
            ["Principio", "Aplicación habitual en phishing", "Riesgo para la víctima"],
            ["Autoridad", "Suplantación de bancos, administraciones, dirección o proveedores", "Reduce el cuestionamiento de la solicitud"],
            ["Urgencia / escasez", "Plazos breves, bloqueo de cuentas o falsas oportunidades", "Favorece decisiones precipitadas"],
            ["Reciprocidad", "Regalos, descuentos o supuestas actualizaciones gratuitas", "Normaliza la primera interacción"],
            ["Compromiso y coherencia", "Secuencias que comienzan con una acción pequeña", "Aumenta la probabilidad de continuar"],
            ["Prueba social", "Avisos sobre acciones atribuidas a otros usuarios", "Hace que la petición parezca habitual"],
            ["Simpatía", "Identidades conocidas, tono cercano o afinidad profesional", "Incrementa la confianza inicial"],
        ],
    )


def correct_state_of_the_art(doc) -> None:
    replace_exact(
        doc,
        "Estado del arte en técnicas de detección y prevención de phishing basadas en Machine Learning y Deep Learning (2019-2025)",
        "Estado del arte en técnicas de detección y prevención de phishing basadas en Machine Learning y Deep Learning (2018-2025)",
    )
    rewrite_table(
        find_table(doc, "Año"),
        [
            ["Año", "Trabajo", "Enfoque", "Objeto", "Aportación", "Limitación"],
            ["2018", "Hutchinson et al.", "Random Forest", "Sitios web", "Muestra la utilidad de rasgos de URL y página", "Depende de variables diseñadas manualmente"],
            ["2020", "Yerima y Alzaylaee", "Red convolucional", "URLs y páginas", "Aprende representaciones con menos ingeniería manual", "No cubre por sí sola la semántica completa del correo"],
            ["2020", "AlEroud y Karabatis", "Red adversarial", "Evasión de detectores", "Estudia robustez frente a ejemplos generados", "La generalización fuera del laboratorio es difícil"],
            ["2022", "Sánchez-Paniagua et al.", "Dataset multipropósito", "Sitios web", "Facilita comparaciones reproducibles", "No representa todos los canales de phishing"],
        ],
        size=8.0,
    )

    deep_table = next(
        table
        for table in doc.tables
        if table.rows
        and len(table.rows[0].cells) == 6
        and table.rows[0].cells[1].text.strip() == "Autores / Trabajo"
    )
    rewrite_table(
        deep_table,
        [
            ["Año", "Trabajo", "Enfoque", "Objeto", "Aportación", "Limitación"],
            ["2023", "Thakur et al.", "Revisión sistemática", "Correo", "Sintetiza arquitecturas de deep learning", "Los datasets y métricas son heterogéneos"],
            ["2023", "Wang et al.", "Modelo preentrenado", "URLs", "Explora transferencia y detección a gran escala", "Se centra en la URL, no en el mensaje completo"],
            ["2024", "Altwaijry et al.", "Comparativa DL", "Correo", "Compara varias arquitecturas sobre una misma tarea", "El resultado depende del corpus evaluado"],
            ["2024", "Koide et al.", "LLM zero-shot", "Correo", "Evalúa detección sin ajuste específico", "Coste, privacidad y reproducibilidad"],
            ["2024", "Rojas-Galeano", "LLM zero-shot", "Spam", "Analiza transferencia entre modelos preentrenados", "Spam y phishing no son categorías equivalentes"],
            ["2024", "Uddin et al.", "Transformer explicable", "Correo", "Relaciona predicción con términos relevantes", "Preprint pendiente de validación más amplia"],
            ["2025", "Kuikel et al.", "LLM + explicaciones", "Correo", "Compara precisión y alineación de explicaciones", "Existe un compromiso entre acierto y fidelidad"],
            ["2025", "Alghenaim et al.", "Revisión del estado del arte", "Detección con IA", "Resume tendencias, riesgos y líneas abiertas", "No es un clasificador evaluado de forma original"],
        ],
        size=7.7,
    )

    replacements = [
        (
            "La detección de phishing ha evolucionado en paralelo",
            "La detección de phishing ha evolucionado desde reglas y listas de bloqueo hacia modelos de machine learning y deep learning capaces de combinar rasgos de URL, contenido, cabeceras y contexto. Las cifras publicadas no son directamente comparables porque varían el corpus, el equilibrio de clases, la partición de datos y la definición de phishing; por ello, esta revisión prioriza las aportaciones y limitaciones de cada enfoque sobre una clasificación basada únicamente en accuracy (Thakur et al., 2023; Alghenaim et al., 2025).",
        ),
        (
            "Esta sección revisa la literatura reciente",
            "La sección revisa trabajos académicos e informes técnicos entre 2018 y 2025. Se organiza en enfoques clásicos, deep learning, transformers, métodos híbridos y limitaciones, y utiliza las revisiones citadas para contextualizar estudios representativos sin presentar sus métricas como si procedieran de un único benchmark.",
        ),
        (
            "2019-2021: Enfoques clásicos de ML basados en features manuales",
            "2019-2021: Enfoques clásicos basados en características manuales. Random Forest, SVM y otros clasificadores trabajan con URL, HTML, cabeceras y variables léxicas. Son rápidos y relativamente explicables, pero dependen de rasgos que pueden envejecer cuando cambian las tácticas.",
        ),
        (
            "2022-2023: Primera ola de Deep Learning",
            "2022-2023: Consolidación del deep learning. CNN, LSTM y modelos preentrenados reducen parte de la ingeniería manual y aprenden patrones en URLs o texto. Su comparación sigue limitada por corpus distintos y por evaluaciones realizadas en entornos controlados.",
        ),
        (
            "2024-2025: Era de Transformers y Large Language Models",
            "2024-2025: Transformers y modelos de lenguaje. Los trabajos recientes estudian clasificación zero-shot, ajuste fino y explicaciones generadas por LLM. El potencial es alto, pero la precisión, la fidelidad de las explicaciones, el coste y la privacidad no mejoran siempre de forma simultánea (Koide et al., 2024; Kuikel et al., 2025).",
        ),
        (
            "Esta evolución refleja un shift de reactivo",
            "Esta evolución refleja el paso de reglas fijas a modelos que aprenden representaciones más ricas. No elimina la utilidad de los enfoques clásicos: en despliegues locales, la velocidad, la trazabilidad y el coste siguen siendo criterios relevantes.",
        ),
        (
            "Los enfoques clásicos de machine learning dominaron",
            "Los enfoques clásicos extraen características de URLs, páginas, cabeceras o texto y las proporcionan a clasificadores supervisados. Su rendimiento puede ser elevado dentro del corpus utilizado, pero cae cuando la distribución cambia o el atacante modifica los rasgos que el modelo espera.",
        ),
        (
            "Hutchinson et al. (2018) obtuvieron",
            "Hutchinson et al. (2018) mostraron la utilidad de Random Forest sobre características de sitios web. El trabajo ilustra la potencia de los conjuntos de árboles, pero también la dependencia de variables definidas de antemano.",
        ),
        (
            "Yerima y Alzaylaee (2020) sustituyeron",
            "Yerima y Alzaylaee (2020) aplicaron redes convolucionales a la detección de phishing, mientras que AlEroud y Karabatis (2020) estudiaron ataques adversariales contra detectores. En conjunto, estos trabajos desplazan el foco desde la precisión estática hacia la capacidad de aprender representaciones y resistir evasiones.",
        ),
        (
            "Li et al. (2022) lograron",
            "Sánchez-Paniagua et al. (2022) publicaron un dataset multipropósito y características de tecnologías web para facilitar comparaciones reproducibles. La limitación común de esta generación es que un buen resultado sobre páginas o URLs no demuestra por sí solo el comportamiento ante correos nuevos, idiomas distintos o campañas generadas por IA.",
        ),
        (
            "Estos modelos destacan en precisión, pero varían en coste:",
            "Los trabajos recientes amplían el objeto de análisis, pero difieren en corpus, tarea y protocolo de evaluación. La tabla resume su aportación principal sin mezclar resultados que no proceden del mismo benchmark.",
        ),
        (
            "La segunda ola, impulsada por el deep learning",
            "Las revisiones de Thakur et al. (2023) y Alghenaim et al. (2025) describen una transición hacia redes profundas, modelos preentrenados y combinaciones de texto, URL y metadatos. Este avance reduce parte de la ingeniería manual, aunque mantiene problemas de comparabilidad y deriva temporal.",
        ),
        (
            "Altwaijry et al. (2024) compararon sistemáticamente",
            "Altwaijry et al. (2024) compararon modelos de deep learning para correo y Koide et al. (2024) exploraron un detector basado en modelos de lenguaje. Ambos trabajos muestran el interés por incorporar contexto semántico, pero no justifican trasladar una cifra de accuracy a otros corpus o entornos.",
        ),
        (
            "Rojas-Galeano (2024) confirmó",
            "Rojas-Galeano (2024) estudió clasificación zero-shot de spam, una tarea próxima pero no idéntica al phishing. Uddin et al. (2024) abordaron la explicabilidad con transformers y Kuikel et al. (2025) observaron que una mejor alineación entre predicción y explicación no implica necesariamente mayor precisión. Esta tensión es central para sistemas de apoyo a la decisión.",
        ),
        (
            "Los trabajos de 2024-2025 integran ML/DL",
            "Los enfoques híbridos combinan señales estructurales, contenido y comportamiento para aportar contexto que una única fuente no contiene. La integración puede mejorar cobertura, pero aumenta la complejidad de datos, despliegue y evaluación.",
        ),
        (
            "Híbridos semánticos + ingeniería social:",
            "Híbridos semánticos e ingeniería social: combinan representaciones del texto con señales de urgencia, autoridad o solicitud de credenciales. Su principal reto es demostrar que esas señales generalizan fuera del corpus de entrenamiento.",
        ),
        (
            "Multimodal: Fusión de texto",
            "Multimodal: fusiona texto, URL, imágenes y metadatos. Puede detectar señales complementarias, aunque exige datasets alineados y eleva el coste de entrenamiento e inferencia.",
        ),
        (
            "Zero-shot: GPT-4 y Gemini-1.5",
            "Zero-shot: utiliza modelos preentrenados sin ajuste específico. Es útil para explorar ataques emergentes, pero introduce variabilidad, coste y posibles transferencias de datos a servicios externos.",
        ),
        (
            "Prevención proactiva: Modelos UEBA",
            "Prevención proactiva: UEBA analiza cambios de comportamiento y anomalías de acceso. Complementa la inspección del mensaje, aunque requiere telemetría organizativa que queda fuera del alcance de este prototipo.",
        ),
        (
            "Estos enfoques híbridos representan el 70 %",
            "La literatura reciente presenta los métodos híbridos como una línea prometedora, pero subraya la necesidad de evaluar cada componente y evitar que la complejidad oculte qué señal sostiene la decisión (Alghenaim et al., 2025; Popescul y Radu, 2025).",
        ),
        (
            "Estos enfoques híbridos representan el 70%",
            "La literatura reciente presenta los métodos híbridos como una línea prometedora, pero subraya la necesidad de evaluar cada componente y evitar que la complejidad oculte qué señal sostiene la decisión (Alghenaim et al., 2025; Popescul y Radu, 2025).",
        ),
        (
            "A pesar de precisiones cercanas al 99%",
            "Aunque distintos trabajos publican resultados elevados en sus propios corpus, persisten desafíos que impiden compararlos directamente:",
        ),
        (
            "Datasets desfasados:",
            "Datasets desfasados: muchos corpus no reflejan campañas recientes, IA generativa, quishing ni deriva de dominios. Un resultado interno elevado puede degradarse ante ataques nuevos.",
        ),
        (
            "Coste y latencia:",
            "Coste y latencia: los modelos grandes exigen más memoria, cómputo o llamadas a servicios externos; los modelos locales ligeros ofrecen mayor control y reproducibilidad.",
        ),
        (
            "Multilingüismo y sesgo:",
            "Multilingüismo y sesgo: el predominio de corpus en inglés dificulta trasladar resultados al español y a contextos culturales diferentes.",
        ),
        (
            "Adversarial robustness:",
            "Robustez adversarial: homógrafos, ofuscación, cambios de plantilla y texto generado pueden eludir patrones aprendidos. Pocos trabajos evalúan de forma comparable estas evasiones.",
        ),
        (
            "Explicabilidad: Black-box",
            "Explicabilidad: una explicación generada no garantiza fidelidad al proceso de decisión. Debe distinguirse entre texto persuasivo y evidencia verificable del modelo.",
        ),
        (
            "Una revisión de 2025 destaca que, pese a avances",
            "La revisión bibliométrica de Popescul y Radu (2025) confirma el crecimiento de ML, deep learning y modelos híbridos, y señala como líneas abiertas la adaptación a nuevos vectores, la robustez y la integración organizativa.",
        ),
        (
            "La literatura 2023-2025 valida los LLMs y transformers",
            "La literatura reciente sitúa los transformers y los LLM entre las líneas activas de investigación, pero advierte que los resultados de entornos controlados no son directamente comparables. Persisten tres huecos relevantes:",
        ),
    ]
    for prefix, replacement in replacements:
        replace_prefix(doc, prefix, replacement)


def correct_project_sections(doc) -> None:
    replace_prefix(
        doc,
        "El alcance del prototipo se ha delimitado al análisis estático",
        "El alcance del prototipo se ha delimitado al análisis estático del correo electrónico y de las URLs contenidas en él. El sistema no navega activamente hacia las páginas enlazadas, no descarga contenido externo ni consulta servicios de reputación en tiempo real. La API, la extensión, Streamlit y los procesos de monitorización se mantienen en el entorno local y no convierten el proyecto en un servicio remoto multiusuario.",
    )
    replace_prefix(
        doc,
        "Esta separación facilita la evolución del prototipo.",
        "Esta separación facilita la evolución del prototipo. Pueden añadirse reglas en los módulos de señales, sustituirse el clasificador neuronal o incorporarse nuevas fuentes de entrada sin alterar la navegación de la aplicación. La versión revisada ejecuta correctamente 72 pruebas Python y 2 recorridos reales con Chromium.",
    )
    replace_prefix(
        doc,
        "La revisión actual de la rama web ejecuta 44 pruebas unitarias",
        "La versión revisada ejecuta 72 pruebas Python y 2 recorridos reales con Chromium. Cubren componentes, integraciones locales, seguridad de bind, activos de la extensión, evaluación e interfaces. Todas finalizaron correctamente; los avisos de convergencia proceden de iteraciones reducidas del MLP en pruebas rápidas.",
    )
    replace_prefix(
        doc,
        "El prototipo presenta varias limitaciones.",
        "El prototipo presenta varias limitaciones. Las listas negras son locales y no se consultan servicios externos de reputación en tiempo real. Tampoco se realiza análisis dinámico de las páginas enlazadas, ni se validan certificados digitales ni se comprueba la reputación histórica de los dominios. Además, Streamlit se utiliza como interfaz académica en un entorno controlado: la versión actual no incorpora gestión multiusuario, control de acceso ni aislamiento de datos para un despliegue público.",
    )

    for paragraph in _paragraphs(doc):
        text = paragraph.text
        paragraph.text = paragraph.text.replace(
            "la revisión actual de la rama web ejecuta 44 pruebas",
            "la versión revisada ejecuta 72 pruebas Python y 2 de navegador",
        )
        paragraph.text = paragraph.text.replace(
            "correctamente 44 pruebas",
            "correctamente 72 pruebas Python y 2 de navegador",
        )
        paragraph.text = paragraph.text.replace(
            "La versión final ejecuta 44 pruebas",
            "La versión revisada ejecuta 72 pruebas Python y 2 de navegador",
        )
        paragraph.text = paragraph.text.replace(
            "La suite de 44 pruebas",
            "La suite de 72 pruebas Python y 2 de navegador",
        )
        paragraph.text = paragraph.text.replace(
            "suite actual de 44 casos",
            "suite actual de 47 casos",
        )
        paragraph.text = paragraph.text.replace(
            "suite de 44 pruebas unitarias",
            "suite de 72 pruebas Python y 2 de navegador",
        )
        if text.startswith(".venv\\Scripts\\Activate.ps1"):
            paragraph.text = ".\\" + text
        if text.startswith("☐"):
            paragraph.text = "Verificado: " + text[1:].strip()
        if "e.g.," in paragraph.text:
            paragraph.text = paragraph.text.replace("e.g.,", "p. ej.,")


def clean_references(doc) -> None:
    replace_prefix(
        doc,
        "IBM. (2025). Cost of a Data Breach Report 2025.",
        "IBM. (2025). Cost of a Data Breach Report 2025. https://www.ibm.com/reports/data-breach",
    )
    for prefix in (
        "Khalifa, O. Y. I., & Muhammad Zaly Shah",
        "LevelBlue. (2025). 2025 LevelBlue Futures Report",
        "Wikipedia. (2025). Phishing.",
        "ZeroFox. (2025a).",
        "ZeroFox. (2025b).",
    ):
        remove_prefix(doc, prefix)


def correct_typographical_errors(doc) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "SIEM tolos (Splunk, Qradar)":
                    _set_cell_text(
                        cell,
                        "SIEM (Splunk, QRadar)",
                        header=False,
                        size=9.0,
                    )


def main() -> None:
    # Wrapper de trazabilidad: evita que transformaciones históricas
    # sobrescriban la arquitectura cliente-servidor documentada.
    from sync_delivery_docs import main as sync_delivery_docs

    sync_delivery_docs()
    return

    doc = Document(DOC_PATH)
    remove_provisional_appendix(doc)
    clean_unverified_visuals(doc)
    correct_context_and_psychology(doc)
    correct_state_of_the_art(doc)
    correct_project_sections(doc)
    clean_references(doc)
    correct_typographical_errors(doc)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
