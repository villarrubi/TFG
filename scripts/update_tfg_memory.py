"""Migración histórica; delega la entrega actual en ``sync_delivery_docs.py``."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BLUE = "2E74B5"
DARK = "1F4D78"
TEXT = "1F2937"
FILL = "E8EEF5"


def set_cell(cell, text, *, bold=False, fill=None):
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(DARK if bold else TEXT)
    run.bold = bold


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_pr.append(table_width)
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    table_pr.append(table_indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = OxmlElement("w:tcW")
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            cell._tc.get_or_add_tcPr().append(tc_width)


def repeat_header(row):
    """Marca la primera fila para lectores de pantalla y páginas repetidas."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def improve_accessibility(doc):
    """Completa metadatos básicos sin alterar el contenido de la memoria."""
    for table in doc.tables:
        if table.rows:
            repeat_header(table.rows[0])
    for index, doc_pr in enumerate(doc.part.element.xpath(".//wp:docPr"), start=1):
        if not doc_pr.get("descr"):
            doc_pr.set("descr", f"Figura o captura del sistema de detección, imagen {index}.")
        if not doc_pr.get("title"):
            doc_pr.set("title", f"Figura del TFG {index}")


def add_table(doc, headers, rows):
    widths = [2100, 7260] if len(headers) == 2 else [1800, 3000, 4560]
    table = doc.add_table(rows=1, cols=len(headers))
    geometry(table, widths)
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, fill=FILL)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], str(value))
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_para(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def main():
    # Wrapper conservado para comandos antiguos. Toda edición documental pasa
    # por el sincronizador cliente-servidor actual.
    from sync_delivery_docs import main as sync_delivery_docs

    sync_delivery_docs()
    return

    path = ROOT / "TFG.docx"
    doc = Document(path)
    replacements = {
        "33 pruebas unitarias y de integración": "41 pruebas unitarias y de integración",
        "33 pruebas unitarias y de integración mediante unittest": "41 pruebas unitarias y de integración mediante unittest",
        "33 pruebas automatizadas": "41 pruebas automatizadas",
        "41 pruebas unitarias y de integración": "89 pruebas unitarias y de integración",
        "41 pruebas unitarias y de integración mediante unittest": "89 pruebas unitarias y de integración mediante unittest",
        "41 pruebas automatizadas": "89 pruebas automatizadas",
        "contiene 41 pruebas.": "contiene 89 pruebas.",
        "la interfaz Streamlit invoca directamente el motor heurístico y el clasificador neuronal dentro del mismo entorno de ejecución": "la interfaz Streamlit actúa como cliente de un backend HTTP obligatorio que centraliza análisis, entrenamiento y modelos",
        "La solución adopta una arquitectura modular autocontenida dentro de una única aplicación Python. Streamlit proporciona la capa de presentación y organiza las vistas de inicio, configuración, detección, monitorización y entrenamiento. Estas vistas invocan directamente los servicios del paquete sistema_phishing en el mismo proceso, por lo que el análisis no depende de un servicio HTTP separado ni requiere enviar el contenido del correo a un servidor central.": "La solución adopta una arquitectura cliente-servidor. Streamlit proporciona la capa de presentación y consume un backend HTTP obligatorio; extensión y monitor usan el mismo contrato y ninguna interfaz carga modelos.",
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)

    improve_accessibility(doc)

    marker = "Actualización técnica y evidencias de la versión final"
    if any(marker in paragraph.text for paragraph in doc.paragraphs):
        doc.save(path)
        return

    doc.add_page_break()
    add_heading(doc, marker)
    add_para(doc, "Esta sección se incorpora para alinear la memoria con la implementación actual de la rama web y dejar trazable qué mejoras se han verificado después de la redacción inicial.")
    add_heading(doc, "Cambios funcionales y de calidad", 2)
    add_table(doc, ["Área", "Implementación verificada"], [
        ("Parseo EML", "Se validan tipo, tamaño y contenido; se conservan cabeceras repetidas y se incluyen todas las cabeceras originales en el texto analizable."),
        ("URLs y remitentes", "Se usan dominios exactos o subdominios válidos, se detectan redirecciones codificadas y se evita marcar como engañoso un nombre personal legítimo."),
        ("Entrenamiento", "Se validan hiperparámetros y etiquetas, cada ejecución parte de los CSV seleccionados y los nuevos joblib no serializan textos brutos."),
        ("Idioma", "El detector selecciona y cachea un modelo por mensaje/idioma, evitando que el primer correo de una sesión fije el idioma de todos."),
        ("Monitor", "Se reutiliza el servicio por lote, se escribe el estado de forma atómica y un mensaje defectuoso no interrumpe los demás."),
        ("API", "Existe un núcleo HTTP común para backend y extensión con límite de 1 MiB, Content-Type obligatorio, CORS restringido y errores sin trazas internas."),
        ("Evaluación", "La aplicación presenta accuracy, precisión, recall, F1, accuracy balanceada y matriz VP/VN/FP/FN sobre un CSV de prueba separado."),
    ])
    add_heading(doc, "Evidencia reproducible", 2)
    add_para(doc, "La suite se ejecuta con python -m unittest discover -s tests -p \"test_*.py\" y contiene 89 pruebas. También se ejecuta python -m ruff check src tests scripts browser_tests, que finaliza sin errores estáticos. Las advertencias de convergencia observadas pertenecen únicamente a pruebas deliberadamente configuradas con 20 iteraciones para comprobar el flujo rápido.")
    add_heading(doc, "Requisitos funcionales explícitos", 2)
    add_table(doc, ["ID", "Requisito y evidencia"], [
        ("RF-01", "Analizar texto pegado, EML y mensajes Gmail sin ejecutar adjuntos."),
        ("RF-02", "Detectar señales de cabeceras, URLs, HTML, contenido y adjuntos y explicar las activas."),
        ("RF-03", "Clasificar con modo heurístico, neuronal o combinado y aplicar un umbral configurable."),
        ("RF-04", "Entrenar modelos español/inglés con CSV y evaluar con un conjunto de prueba independiente."),
        ("RF-05", "Monitorizar Gmail y notificar por Telegram sin bloquear el lote ante un correo defectuoso."),
        ("RF-06", "Ofrecer web, extensión Gmail y monitor como clientes del backend central obligatorio."),
    ])
    add_heading(doc, "Alineación con el alcance", 2)
    add_para(doc, "La versión final mantiene el análisis estático y local. No consulta reputación online, no valida certificados TLS del destino, no navega a las URLs y no presenta la extensión como producto publicado: se carga en modo desarrollador y consulta un servidor local. Estas decisiones se documentan como límites de seguridad y como líneas futuras, no como funcionalidades ya implementadas.")
    add_heading(doc, "Mejoras recomendadas para la defensa", 2)
    for text in [
        "Incluir en la presentación un diagrama de casos de uso y un diagrama de componentes con Streamlit, API/extensión, analysis_service, señales y modelo.",
        "Mostrar una tabla de requisitos RF-01 a RF-06 y relacionarla con pruebas concretas.",
        "Separar siempre métricas de entrenamiento de métricas del conjunto de prueba y explicar VP, VN, FP y FN.",
        "Defender conscientemente el alcance: privacidad local y reproducibilidad frente a reputación online y modelos más pesados.",
    ]:
        paragraph = doc.add_paragraph(style="Lista TFG")
        paragraph.add_run(text)
    doc.save(path)


if __name__ == "__main__":
    main()
