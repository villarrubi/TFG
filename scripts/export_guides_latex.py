"""Exporta las dos guías de defensa DOCX a documentos LaTeX independientes.

Los DOCX siguen siendo la fuente editorial de las guías. Este script genera
una versión ``.tex`` autocontenida para cada guía, conservando títulos, listas,
tablas y los avisos de defensa.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]


def normalize(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def latex_text(text: str) -> str:
    """Escapa texto normal para que también funcione con pdfLaTeX."""

    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in text)


def render_runs(paragraph: Paragraph) -> str:
    """Conserva negrita y cursiva de los párrafos de la guía."""

    if not paragraph.runs:
        return latex_text(normalize(paragraph.text))

    rendered: list[str] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        value = latex_text(run.text)
        if run.bold:
            value = rf"\textbf{{{value}}}"
        if run.italic:
            value = rf"\emph{{{value}}}"
        rendered.append(value)
    return "".join(rendered) or latex_text(normalize(paragraph.text))


def heading_title(text: str) -> str:
    """Quita la numeración escrita en Word y deja que LaTeX la numere."""

    return re.sub(r"^\d+\.\s*", "", normalize(text))


def table_rows(table: Table) -> list[list[str]]:
    return [[normalize(cell.text) for cell in row.cells] for row in table.rows]


def emit_callout(rows: list[list[str]]) -> list[str]:
    text = rows[0][0]
    label, separator, body = text.partition(":")
    if separator:
        content = rf"\textbf{{{latex_text(label + ':')}}} {latex_text(body.strip())}"
    else:
        content = latex_text(text)
    return [
        r"\begin{tcolorbox}[colback=uemclight,colframe=uemcgreen,title=Nota para la defensa]",
        content,
        r"\end{tcolorbox}",
        "",
    ]


def emit_table(rows: list[list[str]], index: int, caption: str) -> list[str]:
    if len(rows) == 1 and len(rows[0]) == 1:
        return emit_callout(rows)

    columns = len(rows[0])
    if columns == 4:
        layout = r"|>{\raggedright\arraybackslash}p{2.45cm}|X|X|X|"
    elif columns == 3:
        layout = r"|>{\raggedright\arraybackslash}p{2.75cm}|X|X|"
    else:
        layout = "|" + "|".join("X" for _ in range(columns)) + "|"

    lines = [
        rf"\begin{{table}}[H]",
        r"  \centering",
        r"  \small",
        rf"  \caption{{{latex_text(caption)}}}",
        rf"  \label{{tab:comparativa-{index}}}",
        rf"  \begin{{tabularx}}{{\textwidth}}{{{layout}}}",
        r"    \hline",
        r"    \rowcolor{uemcgreen}",
    ]
    header = [rf"\textcolor{{white}}{{\textbf{{{latex_text(value)}}}}}" for value in rows[0]]
    lines.extend(["    " + " & ".join(header) + r" \\", r"    \hline"])
    for row_index, row in enumerate(rows[1:], start=1):
        values = [latex_text(value) for value in row]
        if row_index % 2 == 0:
            lines.append(r"    \rowcolor{tablelight}")
        lines.extend(["    " + " & ".join(values) + r" \\", r"    \hline"])
    lines.extend([r"  \end{tabularx}", r"\end{table}", ""])
    return lines


def preamble(title: str, subtitle: str, guide_label: str) -> list[str]:
    return [
        "% Documento generado desde la guía DOCX equivalente.",
        "% Compilación recomendada: pdflatex Guia_XX_*.tex (dos pasadas).",
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[spanish,es-nodecimaldot]{babel}",
        r"\usepackage[a4paper,margin=2.35cm]{geometry}",
        r"\usepackage{lmodern}",
        r"\usepackage{microtype}",
        r"\usepackage[table]{xcolor}",
        r"\usepackage{array}",
        r"\usepackage{tabularx}",
        r"\usepackage{float}",
        r"\usepackage{enumitem}",
        r"\usepackage[most]{tcolorbox}",
        r"\usepackage{fancyhdr}",
        r"\usepackage{hyperref}",
        r"\definecolor{uemcgreen}{HTML}{004C3F}",
        r"\definecolor{uemcgold}{HTML}{E5B93F}",
        r"\definecolor{tfgblue}{HTML}{20566B}",
        r"\definecolor{tablelight}{HTML}{E7F1EE}",
        r"\definecolor{uemclight}{HTML}{F2F7F5}",
        r"\hypersetup{colorlinks=true,linkcolor=uemcgreen,urlcolor=tfgblue,pdftitle={" + latex_text(title) + r"},pdfauthor={Alejandro Villarrubia García}}",
        r"\setlength{\parindent}{0pt}",
        r"\setlength{\parskip}{6pt}",
        r"\renewcommand{\arraystretch}{1.16}",
        r"\setlist[itemize]{leftmargin=1.2cm,itemsep=2pt,topsep=2pt}",
        r"\setlist[enumerate]{leftmargin=1.2cm,itemsep=4pt,topsep=2pt}",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        rf"\fancyhead[L]{{\small\color{{uemcgreen}}{latex_text(guide_label)}}}",
        r"\fancyhead[R]{\small\color{uemcgreen}TFG · UEMC}",
        r"\fancyfoot[C]{\color{uemcgreen}\thepage}",
        r"\renewcommand{\headrulewidth}{0.4pt}",
        r"\renewcommand{\headrule}{\hbox to\headwidth{\color{uemcgold}\leaders\hrule height \headrulewidth\hfill}}",
        r"\begin{document}",
        r"\begin{titlepage}",
        r"  \centering",
        r"  \vspace*{2.4cm}",
        rf"  {{\Large\color{{uemcgreen}}\textbf{{{latex_text(guide_label)}}}\par}}",
        r"  \vspace{1.5cm}",
        rf"  {{\Huge\bfseries\color{{uemcgreen}}{latex_text(title)}\par}}",
        r"  \vspace{0.7cm}",
        rf"  {{\large\color{{tfgblue}}{latex_text(subtitle)}\par}}",
        r"  \vfill",
        r"  {\large Proyecto TFG · versión alineada con el código actual\par}",
        r"  \vspace{0.35cm}",
        r"  {\large Alejandro Villarrubia García\par}",
        r"  \vspace{0.35cm}",
        r"  {\large Universidad Europea Miguel de Cervantes\par}",
        r"\end{titlepage}",
        r"\tableofcontents",
        r"\clearpage",
        "",
    ]


def export(docx_path: Path, output_path: Path) -> None:
    document = Document(docx_path)
    paragraphs_by_id = {paragraph._p: paragraph for paragraph in document.paragraphs}
    lines: list[str] = []

    title = ""
    subtitle = ""
    guide_label = ""
    for paragraph in document.paragraphs:
        text = normalize(paragraph.text)
        if paragraph.style.name == "Title":
            title = text
        elif paragraph.style.name == "Subtitle":
            subtitle = text
        elif text.startswith("GUÍA "):
            guide_label = text

    lines.extend(preamble(title, subtitle, guide_label))
    list_environment: str | None = None
    table_index = 0
    current_section = ""

    def close_list() -> None:
        nonlocal list_environment
        if list_environment:
            lines.append(rf"\end{{{list_environment}}}")
            lines.append("")
            list_environment = None

    for child in document.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = paragraphs_by_id.get(child)
            if paragraph is None:
                continue
            text = normalize(paragraph.text)
            if (
                not text
                or paragraph.style.name in {"Title", "Subtitle"}
                or text.startswith("GUÍA ")
                or text.startswith("Proyecto TFG")
            ):
                continue

            if paragraph.style.name.startswith("Heading"):
                close_list()
                current_section = heading_title(text)
                lines.extend([rf"\section{{{latex_text(current_section)}}}", ""])
                continue

            is_numbered = paragraph.style.name == "List Number"
            is_bulleted = paragraph.style.name == "List Bullet"
            is_manual_numbered = bool(re.match(r"^\d+\.\s+", text))
            if is_numbered or is_bulleted or is_manual_numbered:
                environment = "enumerate" if is_numbered or is_manual_numbered else "itemize"
                if list_environment != environment:
                    close_list()
                    lines.append(rf"\begin{{{environment}}}")
                    list_environment = environment
                item_text = re.sub(r"^\d+\.\s+", "", text) if is_manual_numbered else render_runs(paragraph)
                lines.append("  " + rf"\item {latex_text(item_text) if is_manual_numbered else item_text}")
                continue

            close_list()
            lines.extend([render_runs(paragraph), ""])
        elif tag == "tbl":
            close_list()
            table_index += 1
            if current_section == "Piezas y responsabilidades":
                caption = "Piezas y responsabilidades"
            elif current_section == "API y extensión":
                caption = "Endpoints y controles de la API"
            elif current_section == "Comparativa de tecnologías":
                caption = "Comparativa de tecnologías"
            elif current_section == "SOLID aplicado":
                caption = "Aplicación de SOLID"
            elif current_section == "Diferencias con la propuesta inicial":
                caption = "Diferencias frente a la propuesta inicial"
            else:
                caption = f"Tabla {table_index}"
            lines.extend(emit_table(table_rows(Table(child, document._body)), table_index, caption))

    close_list()
    lines.extend([r"\end{document}", ""])
    output_path.write_text("\n".join(lines), encoding="utf-8", newline="\n")


def main() -> None:
    guides = (
        "Guia_01_Flujo_y_funcionamiento.docx",
        "Guia_02_Tecnologias_y_decisiones.docx",
    )
    for filename in guides:
        source = ROOT / filename
        output = ROOT / filename.replace(".docx", ".tex")
        export(source, output)
        print(output)


if __name__ == "__main__":
    main()
