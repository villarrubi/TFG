"""Genera las tres guías breves de defensa del TFG.

Las guías son entregables versionados, aunque no forman parte del producto
desplegable. Este script es determinista para regenerarlas al cambiar el flujo
o las decisiones tecnológicas del proyecto.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
TEXT = "1F2937"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _keep_row_together(row):
    """Evita que Word parta una fila entre dos páginas."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_run_font(run, *, size=11, color=TEXT, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _set_style(style, *, size=11, color=TEXT, bold=False, before=0, after=6, line=1.25):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def _add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _setup_document(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _set_style(doc.styles["Normal"], after=6, line=1.25)
    _set_style(doc.styles["Title"], size=30, color="203748", bold=True, after=8, line=1.0)
    _set_style(doc.styles["Subtitle"], size=15, color="2B5163", after=18, line=1.1)
    _set_style(doc.styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.1)
    _set_style(doc.styles["Heading 2"], size=13, color=BLUE, bold=True, before=14, after=7, line=1.1)
    _set_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, bold=True, before=10, after=5, line=1.1)
    _set_style(doc.styles["List Bullet"], after=4, line=1.25)
    _set_style(doc.styles["List Number"], after=4, line=1.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(header.add_run("TFG · Guía de defensa"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer.add_run("Página "), size=9, color=MUTED)
    _add_page_field(footer)
    doc.core_properties.title = title
    doc.core_properties.author = "Proyecto TFG"
    doc.core_properties.subject = "Preparación de defensa del TFG"
    return doc


def _cover(doc, kicker, title, subtitle):
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    _set_run_font(p.add_run(kicker.upper()), size=10, color="B4862C", bold=True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(title), size=30, color="203748", bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(subtitle), size=15, color="2B5163")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    _set_run_font(p.add_run("Proyecto TFG · versión alineada con el código actual"), size=10, color=MUTED, italic=True)
    doc.add_page_break()


def _heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def _para(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        _set_run_font(p.add_run(bold_prefix), bold=True)
        _set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        _set_run_font(p.add_run(text))
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _set_run_font(p.add_run(text))
    return p


def _number(doc, text):
    p = doc.add_paragraph(style="List Number")
    _set_run_font(p.add_run(text))
    return p


def _callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [TABLE_WIDTH])
    _repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    _set_run_font(p.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
    _set_run_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    _set_table_geometry(table, widths)
    header = table.rows[0]
    _repeat_header(header)
    _keep_row_together(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        _set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        _set_run_font(p.add_run(value), bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        _keep_row_together(row)
        cells = row.cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            _set_run_font(p.add_run(str(value)))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def flujo():
    doc = _setup_document("Flujo y funcionamiento · TFG Phishing")
    _cover(doc, "Guía 01 · mapa del sistema", "Flujo y funcionamiento", "Qué ocurre desde que entra un correo hasta que se explica el riesgo")
    _heading(doc, "1. Idea general", 1)
    _para(doc, "El montaje es cliente-servidor. El navegador usa Streamlit como capa de presentación y Streamlit envía peticiones HTTP/JSON al backend central. La extensión y el monitor también consumen esa API. Solo el servidor parsea, analiza, entrena y guarda los modelos; Gmail y Telegram son integraciones externas.")
    _callout(doc, "Frase para la defensa", "Aunque cliente y servidor estén en el mismo ordenador, son procesos separados con un contrato HTTP. Hay una sola versión activa por idioma y todos los clientes la comparten.")
    _heading(doc, "2. Flujo completo", 1)
    for text in [
        "Entrada de cliente: la web, la extensión o el monitor envían texto, campos JSON o un EML Base64 al backend; /analyze limita la petición a 16 MiB.",
        "Normalización: se extraen remitente, asunto, cuerpo, HTML, cabeceras completas, URLs, anclas y adjuntos.",
        "Señales: se revisan autenticación, coherencia de dominios, URLs, punycode, acortadores, HTML, formularios, lenguaje urgente y adjuntos.",
        "Puntuación heurística: RiskScorer pondera las señales activas y produce un riesgo de 0 a 100; ExplanationBuilder conserva el motivo.",
        "Idioma y modelo: EmailAnalysisService detecta español o inglés por mensaje y reutiliza un detector neuronal cacheado por idioma.",
        "Decisión: el modo combinado calcula una media ponderada; el umbral configurado se aplica de forma uniforme en heurístico, neuronal y combinado.",
        "Salida: el backend devuelve clasificación, puntuación, explicación, señales, idioma y versión; el cliente se limita a presentarlos o generar la alerta.",
    ]:
        _number(doc, text)
    _heading(doc, "3. Piezas y responsabilidades", 1)
    _table(doc, ["Capa", "Módulo", "Responsabilidad", "Resultado"], [
        ("Entrada", "analizador_email.py", "Parseo MIME seguro y extracción de partes", "Correo normalizado"),
        ("Dominio", "correo.py", "Contrato común para texto, diccionario o EML", "CorreoAnalizado"),
        ("Señales", "header_signals, url_utils, html_signals, content_signals", "Reglas pequeñas y comprobables", "Diccionario de señales"),
        ("Cliente", "backend_client.py", "Contrato HTTP común sin lógica de detección", "Petición/respuesta JSON"),
        ("Servidor", "backend_service.py", "Casos de uso, idioma, versiones y administración", "Resultado central"),
        ("ML", "modelo_neural.py", "TF-IDF, MLP, entrenamiento y persistencia del servidor", "Probabilidad phishing"),
        ("Presentación", "Streamlit, monitor, extensión", "Recoger entrada y mostrar o alertar", "UI o alerta"),
    ], [1500, 2350, 3150, 2360])
    _heading(doc, "4. Cómo se analiza un EML", 1)
    _para(doc, "El parser usa email.parser.BytesParser con una política MIME estándar. Conserva todas las cabeceras repetidas, no solo la última, y serializa las cabeceras originales junto al cuerpo para que SPF, DKIM, DMARC, Received, Return-Path y Message-ID sean visibles a las reglas.")
    for text in [
        "Texto plano: se decodifica respetando el charset declarado y se conserva como cuerpo analizable.",
        "HTML: se guarda separado para revisar formularios, meta refresh, iframes, base href y javascript/data URLs.",
        "Adjuntos: se anotan nombre y tipo sin ejecutar ni abrir contenido potencialmente peligroso.",
        "URLs: se deduplican, se limpian de puntuación y se extraen también de enlaces HTML.",
    ]:
        _bullet(doc, text)
    _heading(doc, "5. Monitor y persistencia", 1)
    _para(doc, "El monitor consulta Gmail, carga los IDs vistos y procesa solo mensajes nuevos. Cada correo normalizado se envía al backend con RemoteAnalysisService; el monitor no carga modelos. El estado se escribe después de cada mensaje completado mediante fichero temporal y os.replace.")
    _para(doc, "Si un EML está corrupto o falla una alerta de Telegram, ese mensaje queda reportado con error y el resto del lote continúa. Un fallo de entrada no se convierte en una caída global del proceso.")
    _heading(doc, "6. API y extensión", 1)
    _table(doc, ["Endpoint", "Entrada", "Respuesta", "Controles"], [
        ("GET /health y /models", "Sin cuerpo", "Estado y versiones sin rutas locales", "No-store; sin artefactos"),
        ("POST /analyze", "JSON hasta 16 MiB", "Riesgo, explicación, idioma y versión", "Tipos, tamaño y CORS"),
        ("POST /train y /evaluate", "CSV serializado hasta 256 MiB", "Versión y métricas", "Token admin opcional en local"),
        ("POST /compare y /models/delete", "Configuraciones o idioma", "Comparación o confirmación", "Token y límites"),
    ], [1800, 2500, 2800, 2260])
    _para(doc, "La extensión de Gmail no ejecuta Python: recoge el correo visible y lo envía directamente al backend central, por defecto en 127.0.0.1:8766. El proceso del puerto 8765 se conserva únicamente como proxy de compatibilidad y tampoco carga modelos.")
    _heading(doc, "7. Ejemplo de extremo a extremo", 1)
    for index, text in enumerate([
        "El usuario abre un correo con asunto urgente y un enlace que aparenta ser de una marca conocida.",
        "El adaptador extrae From, Subject, cuerpo, anclas y URL.",
        "Las reglas detectan urgencia, dominio externo, anchor distinto y posible incoherencia de cabeceras.",
        "El modelo neuronal calcula una probabilidad sobre el texto completo.",
        "El modo combinado aplica la fusión calibrada 35/65 y el umbral 26; si una evidencia alcanza 70 no se diluye, y la explicación muestra qué señales activaron el riesgo.",
        "La UI pinta la tarjeta, la extensión la inserta en Gmail y el monitor podría notificar por Telegram.",
    ], 1):
        _para(doc, f"{index}. {text}")
    _heading(doc, "8. Límites que conviene decir", 1)
    _bullet(doc, "La lista de dominios y señales es local; no se afirma que exista reputación online en tiempo real.")
    _bullet(doc, "El análisis es apoyo a la decisión, no sustituto de una pasarela antispam ni garantía absoluta.")
    _bullet(doc, "El MLP necesita datos representativos y evaluación separada; una accuracy de entrenamiento no demuestra generalización.")
    _para(
        doc,
        "Validación: 87 pruebas Python y 2 recorridos reales de navegador; uno "
        "levanta Streamlit y backend por separado y otra prueba local recorre Gmail, HTTP y Telegram. Comprueban comportamiento y "
        "arquitectura, no eficacia estadística en producción.",
    )
    doc.save(ROOT / "Guia_01_Flujo_y_funcionamiento.docx")


def tecnologias():
    doc = _setup_document("Tecnologías y decisiones · TFG Phishing")
    _cover(doc, "Guía 02 · decisiones técnicas", "Tecnologías y decisiones", "Cómo funciona cada elección y cómo justificarla frente a alternativas")
    _heading(doc, "1. Criterio de selección", 1)
    _para(doc, "La arquitectura prioriza reproducibilidad local, facilidad de defensa, separación de responsabilidades y coste operativo bajo. Las decisiones no intentan convertir el TFG en un servicio de producción multiusuario: delimitan un detector demostrable, auditable y extensible.")
    _heading(doc, "2. Comparativa de tecnologías", 1)
    _table(doc, ["Tecnología", "Cómo se usa", "Por qué se elige", "Alternativa y límite"], [
        ("Python", "Integra parser, reglas, ML, API y automatizaciones", "Ecosistema científico, lectura clara y buena trazabilidad", "Java/Node serían válidos; añadirían complejidad para este alcance"),
        ("Streamlit", "Cliente web y capa de presentación", "Permite validar el flujo cliente-servidor con poco código de interfaz", "React daría comunicación navegador-API directa, pero exigiría otro stack"),
        ("scikit-learn", "TF-IDF vectoriza y MLPClassifier clasifica", "Pipeline reproducible, API madura y métricas estándar", "TensorFlow sería más flexible para deep learning, pero sobredimensionado para texto tabular"),
        ("TF-IDF", "Convierte palabras y bigramas en variables", "Explicable, rápido y adecuado para corpus moderados", "Transformers mejorarían contexto, con más coste y menor explicabilidad local"),
        ("email estándar", "Parsea MIME y cabeceras EML", "No añade dependencia para la parte esencial del formato", "Librerías externas pueden aportar atajos, pero no son necesarias"),
        ("BeautifulSoup", "Inspecciona HTML y formularios", "Tolera HTML imperfecto de correos", "Parser regex puro sería frágil; un navegador sería inseguro e innecesario"),
        ("Gmail API + OAuth", "Obtiene mensajes con consentimiento del usuario", "Permisos oficiales y flujo revocable", "IMAP sería más genérico, pero no aporta el mismo control API"),
        ("HTTP stdlib", "Expone análisis y administración central", "Cero framework extra, contrato pequeño y fácil de auditar", "FastAPI facilitaría OpenAPI y escalado; no es necesario para el prototipo"),
        ("joblib", "Persiste pipeline TF-IDF + MLP", "Guarda vectorizador y modelo como una unidad", "ONNX sería más portable; exige conversión y no elimina el requisito de confiar en artefactos"),
        ("unittest + Ruff", "Prueba comportamiento y calidad estática", "Incluidos en Python y rápidos en CI/local", "pytest aporta fixtures más ricas; el conjunto actual es suficiente y directo"),
    ], [1500, 2500, 2800, 2560])
    _heading(doc, "3. Por qué dos modelos lingüísticos", 1)
    _para(doc, "El detector identifica idioma por mensaje, no por sesión. Mantiene un modelo español y otro inglés porque las palabras, stopwords y corpus tienen distribuciones distintas. EmailAnalysisService cachea un detector por idioma para combinar corrección con rendimiento.")
    _callout(doc, "Respuesta breve", "La alternativa de un único modelo multilingüe simplificaría ficheros, pero mezclaría vocabularios y haría más difícil justificar qué datos explican cada predicción.")
    _heading(doc, "4. SOLID aplicado", 1)
    _table(doc, ["Principio", "Aplicación en el código", "Beneficio demostrable"], [
        ("Responsabilidad única", "Parser, señales, scorer, explicaciones, ML, monitor y transporte están separados", "Un cambio en URLs no obliga a tocar Gmail o Streamlit"),
        ("Abierto/cerrado", "EmailAnalysisService recibe analyzer, detector loader y detector de idioma inyectables", "Se prueban estrategias y se añaden modelos sin reescribir el coordinador"),
        ("Sustitución", "Los consumidores dependen de BackendClient o su protocolo mínimo", "Web, extensión y monitor consumen el mismo contrato HTTP"),
        ("Segregación", "El almacenamiento, el análisis y el transporte tienen superficies pequeñas", "Cada prueba necesita solo la interfaz que utiliza"),
        ("Inversión", "NeuralModelTrainer depende de ModelStorage, no de una ruta fija", "Facilita memoria, disco y dobles de prueba"),
    ], [1700, 4300, 3360])
    _heading(doc, "5. Decisiones de seguridad", 1)
    for text in [
        "El parser limita EML a 10 MiB; /analyze admite 16 MiB y las rutas de datasets 256 MiB con Content-Length obligatorio.",
        "La API no acepta CORS *; permite Gmail y orígenes chrome-extension://, y no devuelve trazas internas.",
        "Streamlit y el backend escuchan en loopback; --allow-remote exige un token administrativo de al menos 24 caracteres y los clientes exigen HTTPS para destinos remotos.",
        "El modelo no ejecuta archivos adjuntos; solo registra metadatos y señales.",
        "Los nuevos joblib no conservan textos brutos de entrenamiento. Los modelos antiguos se sanean al cargarse y guardarse con el formato actual.",
        "Credenciales, tokens, estados y artefactos temporales de QA se mantienen fuera del índice Git mediante .gitignore.",
    ]:
        _bullet(doc, text)
    _heading(doc, "6. Reproducibilidad y evaluación", 1)
    _para(doc, "Los hiperparámetros se concentran en HiperparametrosModelo. El cliente los envía con los CSV y el backend entrena desde cero, registra procedencia, guarda atómicamente e invalida la versión cacheada. Cuarenta casos controlados calibran la fusión y 16 EML distintos, con cabeceras y MIME completos, quedan reservados para la evaluación final; ambos conjuntos son sintéticos. Un diagnóstico adicional usa 1.528 textos DIFrauD con licencia MIT, revisión y hash fijados, pero se etiqueta con riesgo de solapamiento de fuentes. Ninguna cifra estima producción.")
    _para(doc, "El benchmark reproducible separa arranque e inferencia. La carga diferida redujo la importación fría de heurísticas de 1098,5 ms a 37,0 ms y el arranque de la aplicación de 1326,6 ms a 315,1 ms, mientras que la inferencia se mantuvo estable dentro de una variación de ±3 %.")
    _para(doc, "La validación automatizada reúne 87 pruebas Python y 2 pruebas con Chromium. Incluye idioma determinista y el recorrido local Gmail EML -> JSON -> HTTP -> backend -> Telegram. GitHub Actions instala versiones fijadas, ejecuta Ruff, verifica calibración, regenera la evaluación, comprueba el respaldo de defensa y recorre web-backend. OAuth real usa una lista E2E separada porque sus credenciales no pueden entrar en CI.")
    _heading(doc, "7. Diferencias con la propuesta inicial", 1)
    _table(doc, ["Propuesta", "Implementación actual", "Cómo explicarlo"], [
        ("TensorFlow", "scikit-learn MLP", "Se eligió una red suficiente para el corpus y más reproducible/ligera en el entorno académico"),
        ("Blacklists/reputación", "Señales locales de dominio, URL, punycode y acortadores", "Se evita depender de red externa y se mantiene privacidad; reputación online queda como extensión"),
        ("Certificados", "No se inspecciona TLS desde el correo", "El EML no prueba por sí mismo el certificado del destino; se prioriza análisis estático seguro"),
        ("Interfaz simple", "Streamlit como cliente de una API central obligatoria", "La interfaz sigue siendo simple y todos los canales comparten los modelos del servidor"),
    ], [2400, 3000, 3960])
    _heading(doc, "8. Qué no prometer", 1)
    _bullet(doc, "No decir que se consulta una blacklist online si no se ha configurado una fuente externa.")
    _bullet(doc, "No presentar el accuracy de entrenamiento como precisión en producción.")
    _bullet(doc, "No llamar a la extensión un producto publicado: se carga en modo desarrollador y consume el backend central configurado.")
    doc.save(ROOT / "Guia_02_Tecnologias_y_decisiones.docx")


def guion():
    doc = _setup_document("Guion de defensa · TFG Phishing")
    _cover(doc, "Guía 03 · exposición oral", "Guion de defensa", "Qué decir, qué enseñar y cómo responder preguntas sin sobreprometer")
    _heading(doc, "1. Marco de tiempo", 1)
    _callout(doc, "Regla práctica", "La regulación pública UEMC del TFG contempla una exposición de hasta 20 minutos y un turno posterior de preguntas; confirma siempre el tiempo exacto indicado por tu centro y convocatoria.")
    _para(doc, "Referencia para preparar la defensa: Reglamento UEMC 5/2024, de 13 de septiembre, de Trabajo Fin de Grado/Fin de Máster, y la ficha pública de la asignatura. El reglamento permite que el centro limite excepcionalmente la exposición a menos de 20 minutos; confirma la convocatoria concreta. La memoria y la defensa deben contar la misma historia: problema, objetivos, método, evidencia y límites.")
    _heading(doc, "2. Estructura recomendada de 20 minutos", 1)
    _table(doc, ["Tiempo", "Bloque", "Mensaje que debe quedar"], [
        ("0:00–1:00", "Apertura", "Qué problema resuelvo y qué entrego"),
        ("1:00–3:00", "Motivación y objetivos", "El phishing combina señales técnicas y lingüísticas; se necesita apoyo explicable"),
        ("3:00–5:00", "Alcance y requisitos", "Texto/EML, heurística, ML, explicación, Gmail, evaluación y límites"),
        ("5:00–9:00", "Arquitectura y flujo", "Entradas -> normalización -> señales/modelo -> decisión -> explicación"),
        ("9:00–12:00", "Tecnologías", "Por qué Python, Streamlit, scikit-learn, Gmail API y backend central"),
        ("12:00–15:00", "Demostración", "Un correo sospechoso, señales activas y resultado combinado"),
        ("15:00–17:00", "Pruebas y resultados", "87 pruebas Python, 2 de navegador, calibración, 16 EML y diagnóstico externo"),
        ("17:00–19:00", "Limitaciones y futuro", "Qué no hace hoy y qué ampliaría con evidencia"),
        ("19:00–20:00", "Cierre", "Aportación, mantenibilidad y conclusión"),
    ], [1500, 2500, 5360])
    _heading(doc, "3. Texto de apertura", 1)
    _para(doc, "Buenos días. Mi TFG presenta un sistema cliente-servidor de apoyo a la detección de phishing. Combina señales inspeccionables —cabeceras, autenticación, dominios, URLs, HTML y lenguaje— con un clasificador TF-IDF más MLP. Devuelve una puntuación y una explicación, no solo una etiqueta.")
    _para(doc, "He separado los clientes del núcleo de análisis. Streamlit, la extensión y el monitor envían el correo a un backend central que mantiene una versión por idioma. Así, al cambiar un modelo, todos los clientes usan la nueva versión sin distribuir copias.")
    _heading(doc, "4. Qué enseñar en la demo", 1)
    for text in [
        "Arrancar el backend, mostrar /health con las versiones y después abrir el cliente web.",
        "Pegar o cargar un EML de ejemplo controlado, evitando enseñar credenciales o datos personales.",
        "Señalar la puntuación, las señales activas, la URL sospechosa y la explicación generada.",
        "Cambiar a modo heurístico para demostrar que el resultado puede auditarse sin modelo.",
        "Mostrar EVALUATION_REPORT.md y EXTERNAL_EVALUATION_REPORT.md: 16 EML, 1.528 textos, hashes, métricas y las advertencias de representatividad/solapamiento.",
        "Si el tiempo lo permite, enseñar la extensión llamando a la misma URL del backend.",
    ]:
        _number(doc, text)
    _callout(doc, "Plan B", "Lleva las capturas enumeradas en docs/DEFENSE_SCREENSHOTS.md y defense_demo/expected_results.json. Si Gmail u OAuth fallan, usa los EML sintéticos; si el backend no arranca, enseña el bloque health y los dos resultados guardados.")
    _heading(doc, "5. Cómo explicar los resultados", 1)
    _para(doc, "La calibración usa 40 casos distintos de los 16 EML finales. Tras añadir tres señales BEC bilingües, la rejilla selecciona 35 % heurística, 65 % neuronal, umbral 26 y alta confianza 70. En los EML, el heurístico obtiene 100,0 % de accuracy/recall/F1; el combinado, 93,8 % de accuracy, 88,9 % de precisión, 100,0 % de recall y 94,1 % de F1; y el neuronal, 75,0 % en las cuatro métricas. DIFrauD añade 1.528 textos externos y el combinado alcanza 90,8 % de accuracy y 96,4 % de recall, pero existe riesgo de solapamiento de fuentes. Ninguna cifra debe presentarse como producción.")
    _para(doc, "En rendimiento, presenta solo mejoras medidas: la importación fría de heurísticas bajó un 96,6 % y el arranque de la aplicación un 76,2 %. Las rutas de inferencia variaron menos de un 3 %, por lo que no se atribuye una mejora que no esté respaldada por la medición.")
    _heading(doc, "6. Preguntas previsibles y respuestas", 1)
    _table(doc, ["Pregunta", "Respuesta breve y defendible"], [
        ("¿Por qué no usar deep learning más grande?", "El corpus y el alcance no lo justificaban. TF-IDF + MLP ofrece rapidez, reproducibilidad y una línea base explicable; la arquitectura permite sustituir el modelo."),
        ("¿Qué ocurre con un correo en inglés?", "El idioma se detecta por mensaje y se cachea un modelo específico por idioma; no se fija el idioma para toda la sesión."),
        ("¿Es cliente-servidor?", "Sí. Streamlit, extensión y monitor son clientes HTTP; backend_server analiza y mantiene una versión activa por idioma. En local ambos lados están en el mismo equipo, pero siguen siendo procesos y responsabilidades separadas."),
        ("¿Cómo evitas falsos positivos?", "La fusión 35/65 y el umbral 26 se calibraron separadamente; conserva evidencias sobre 70, muestra FP/FN y mantiene la explicación para revisar cada caso."),
        ("¿Consulta una blacklist externa?", "No en la versión actual. Hay comprobaciones locales de dominios, URLs, punycode y acortadores; una reputación online sería una ampliación con dependencia y privacidad adicionales."),
        ("¿Es seguro cargar joblib?", "Solo se cargan artefactos propios y verificados, porque joblib usa deserialización Python. No se deben aceptar modelos arbitrarios."),
        ("¿Qué pasa si Gmail está caído?", "La detección por texto o EML sigue funcionando contra el backend; Gmail es solo una fuente externa y muestra un error controlado."),
        ("¿Por qué no guardas todos los correos de entrenamiento?", "Para reducir exposición de datos y hacer el artefacto más limpio. Se conservan procedencia y estadísticas, no el texto bruto."),
        ("¿Qué falta para producción?", "TLS, autenticación de inferencia, rate limiting, registro compartido si hay réplicas, monitorización y confirmación sobre un corpus real reciente, bilingüe e independiente."),
    ], [3300, 6060])
    _heading(doc, "7. Cierre", 1)
    _para(
        doc,
        "Aporta una solución cliente-servidor mantenible: parser seguro, señales "
        "especializadas, modelos centrales, explicación y clientes con un contrato "
        "común. No sustituye a un producto antispam; demuestra un detector "
        "desplegable sin duplicar modelos.",
    )
    _heading(doc, "8. Checklist de la víspera", 1)
    for text in [
        "Ejecutar la suite automatizada completa y guardar la salida.",
        "Comprobar que backend y web arrancan por separado y que /health muestra las versiones.",
        "Regenerar EVALUATION_REPORT.md y explicar sus errores sin confundir el desafío sintético con producción.",
        "Ensayar el guion con cronómetro y dejar dos minutos de margen.",
        "Memorizar tres límites: sin reputación online, extensión local y necesidad de datos representativos.",
        "Llevar capturas del flujo y de la matriz de confusión como respaldo.",
    ]:
        _bullet(doc, text)
    doc.save(ROOT / "Guia_03_Guion_defensa.docx")


if __name__ == "__main__":
    flujo()
    tecnologias()
    guion()
