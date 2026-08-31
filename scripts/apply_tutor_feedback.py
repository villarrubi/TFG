"""Aplica a la memoria y a la guía extensa las observaciones del tutor."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
CONTENT_WIDTH = 8384
BLUE = "12386B"
TEAL = "1D6E7A"
TEXT = "263238"
GRID = "D5DEE8"


def _all_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _set_text(paragraph: Paragraph, value: str) -> None:
    paragraph.clear()
    paragraph.add_run(value)


def _paragraph(doc: DocumentObject, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"No se encontró el párrafo {text!r}")


def _replace_prefix(doc: DocumentObject, prefix: str, replacement: str) -> int:
    changed = 0
    for paragraph in _all_paragraphs(doc):
        if paragraph.text.startswith(prefix):
            _set_text(paragraph, replacement)
            changed += 1
    return changed


def _replace_fragment(doc: DocumentObject, old: str, new: str) -> int:
    changed = 0
    for paragraph in _all_paragraphs(doc):
        if old in paragraph.text:
            _set_text(paragraph, paragraph.text.replace(old, new))
            changed += 1
    return changed


def _remove(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_between(doc: DocumentObject, start: str, end: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_index = next(
        (
            index
            for index, child in enumerate(children)
            if child.tag == qn("w:p") and "".join(child.xpath(".//w:t/text()")) == start
        ),
        None,
    )
    if start_index is None:
        return
    end_index = next(
        (
            index
            for index, child in enumerate(children[start_index + 1 :], start_index + 1)
            if child.tag == qn("w:p") and "".join(child.xpath(".//w:t/text()")) == end
        ),
        None,
    )
    if end_index is None:
        raise ValueError(f"No se encontró el final {end!r} para retirar {start!r}.")
    for child in children[start_index:end_index]:
        _remove(child)


def _insert_paragraph_before(
    doc: DocumentObject,
    anchor: Paragraph,
    text: str,
    *,
    style: str = "Normal",
) -> Paragraph:
    paragraph = doc.add_paragraph(style=style)
    _set_text(paragraph, text)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if paragraph.style is not None:
        inserted.style = paragraph.style
    _set_text(inserted, text)
    return inserted


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcW")
    if node is None:
        node = OxmlElement("w:tcW")
        tc_pr.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH:
        raise ValueError("La geometría de tabla debe sumar el ancho útil.")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_WIDTH))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(cell, value: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value if name in {"top", "bottom"} else 110))
        node.set(qn("w:type"), "dxa")


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _format_table(table: Table, widths: list[int]) -> None:
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    _repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell)
            if row_index == 0:
                _shade(cell, BLUE)
            elif row_index % 2 == 0:
                _shade(cell, "F2F6F9")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.3)
                    run.bold = row_index == 0
                    run.font.color.rgb = RGBColor.from_string(
                        "FFFFFF" if row_index == 0 else TEXT
                    )


def _insert_table_before(
    doc: DocumentObject,
    anchor: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> Table:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    _format_table(table, widths)
    anchor._p.addprevious(table._tbl)
    return table


def _insert_caption(
    doc: DocumentObject,
    anchor: Paragraph,
    text: str,
    *,
    keep_with_next: bool = True,
) -> Paragraph:
    caption = _insert_paragraph_before(doc, anchor, text, style="Pie TFG")
    caption.paragraph_format.keep_with_next = keep_with_next
    caption.paragraph_format.keep_together = True
    return caption


def _insert_picture(
    doc: DocumentObject,
    anchor: Paragraph,
    path: Path,
    *,
    width: float,
    alt_text: str,
) -> Paragraph:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _ensure_front_matter_entries(doc: DocumentObject) -> None:
    figure_entries = [
        "Figura 6.1: Pantalla inicial del cliente web con el backend central conectado.",
        "Figura 6.2: Resultado combinado del escenario BEC sintético reservado.",
    ]
    table_entries = [
        "Tabla 5.1: Cambios respecto a la propuesta inicial.",
        "Tabla 6.1: Conjuntos empleados y separación experimental.",
        "Tabla 6.2: Métricas sobre los 16 EML reservados.",
        "Tabla 6.3: Matrices de confusión sobre los 16 EML reservados.",
    ]
    if not all(
        any(p.text == value for p in doc.paragraphs) for value in figure_entries
    ):
        anchor = _paragraph(
            doc,
            "Figura 2.2: Vectores iniciales en brechas analizadas por el DBIR 2025 (n=9.891).",
        )
        for value in figure_entries:
            if not any(p.text == value for p in doc.paragraphs):
                anchor = _insert_paragraph_after(anchor, value)
    if not all(any(p.text == value for p in doc.paragraphs) for value in table_entries):
        anchor = _paragraph(doc, "Tabla 4.2: Herramientas y tecnologías utilizadas.")
        for value in table_entries:
            if not any(p.text == value for p in doc.paragraphs):
                anchor = _insert_paragraph_after(anchor, value)


def _build_proposal_changes(doc: DocumentObject) -> None:
    marker = "Cambios respecto a la propuesta inicial"
    if any(p.text == marker for p in doc.paragraphs):
        _remove_between(doc, marker, "Arquitectura general")
    anchor = _paragraph(doc, "Arquitectura general")
    _insert_paragraph_before(doc, anchor, marker, style="Heading 2")
    _insert_paragraph_before(
        doc,
        anchor,
        "La propuesta de octubre de 2025 contemplaba un prototipo con reglas, "
        "servicios externos de reputación, comprobación de certificados y la posible "
        "utilización de TensorFlow. Durante el desarrollo se ajustó el alcance para "
        "priorizar privacidad, seguridad de la evaluación, reproducibilidad local y una "
        "arquitectura cliente-servidor con un único modelo compartido.",
    )
    _insert_caption(doc, anchor, "Tabla 5.1: Cambios respecto a la propuesta inicial.")
    _insert_table_before(
        doc,
        anchor,
        ["Elemento", "Propuesta", "Implementación final", "Justificación"],
        [
            [
                "Aprendizaje",
                "scikit-learn y TensorFlow",
                "TF-IDF + MLPClassifier de scikit-learn",
                "Pipeline ligero, reproducible y suficiente para el alcance; TensorFlow no era necesario.",
            ],
            [
                "Reputación",
                "Listas negras y servicios externos",
                "Listas y señales locales; sin consultas online",
                "Evita enviar URLs sensibles, depender de cuotas o visitar infraestructura potencialmente maliciosa.",
            ],
            [
                "Certificados",
                "Comprobación de certificados digitales",
                "Sin conexión al destino ni validación TLS",
                "Un EML no acredita el certificado actual del enlace; la comprobación activa queda como trabajo futuro aislado.",
            ],
            [
                "Arquitectura",
                "Interfaz sencilla del prototipo",
                "Clientes HTTP y backend central obligatorio",
                "Una actualización del modelo se aplica a web, extensión y monitor sin distribuir copias.",
            ],
        ],
        [1200, 1960, 2420, 2804],
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "Estos cambios no alteran el objetivo académico de estudiar, implementar y "
        "evaluar la detección de phishing. Delimitan qué se demuestra en el prototipo "
        "y qué capacidades necesitan infraestructura y controles adicionales.",
    )


def _build_experiment(doc: DocumentObject) -> None:
    start = (
        "Diseño experimental y datasets"
        if any(p.text == "Diseño experimental y datasets" for p in doc.paragraphs)
        else "Dataset de ejemplo"
    )
    _remove_between(doc, start, "Comportamiento del sistema heurístico")
    anchor = _paragraph(doc, "Comportamiento del sistema heurístico")
    _insert_paragraph_before(
        doc, anchor, "Diseño experimental y datasets", style="Heading 2"
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "El protocolo reproducible fija las fuentes, licencias y SHA-256 sin versionar "
        "los mensajes. En español se combinan el split de entrenamiento de "
        "softecapps/spam_ham_spanish (Hugging Face, DOI 10.57967/hf/2264) y SMS Spam "
        "Mexico (Kaggle). Tras retirar 147 copias, dos filas contradictorias y una "
        "coincidencia con la prueba quedan 1.148 textos. Las 209 filas del split oficial "
        "se reservan para probar. En inglés se utiliza únicamente phishing_email.csv "
        "(Kaggle), porque ya agrega CEAS, Enron, Ling, Nazario, Nigerian Fraud y "
        "SpamAssassin: añadir también esos seis CSV duplicaba el corpus. Tras retirar "
        "408 copias quedan 82.077 textos, divididos de forma estratificada 80/20 con "
        "semilla 42. Los joblib guardan el pipeline y estos metadatos, nunca el texto bruto.",
    )
    _insert_caption(
        doc, anchor, "Tabla 6.1: Conjuntos empleados y separación experimental."
    )
    _insert_table_before(
        doc,
        anchor,
        ["Conjunto", "N", "Phishing", "Legítimos", "Procedencia", "Uso"],
        [
            [
                "Entrenamiento ES",
                "1.148",
                "613",
                "535",
                "Hugging Face + SMS Spam Mexico",
                "Ajuste del modelo español",
            ],
            [
                "Prueba ES",
                "209",
                "100",
                "109",
                "Split oficial softecapps",
                "Holdout español",
            ],
            [
                "Entrenamiento EN",
                "65.661",
                "34.275",
                "31.386",
                "80 % del agregado deduplicado",
                "Ajuste del modelo inglés",
            ],
            [
                "Prueba EN",
                "16.416",
                "8.569",
                "7.847",
                "20 % estratificado del agregado",
                "Holdout inglés",
            ],
            [
                "Calibración",
                "40",
                "20",
                "20",
                "Casos sintéticos ES/EN; 10 por idioma y clase",
                "Cinco particiones para pesos y umbral",
            ],
            [
                "Evaluación EML",
                "16",
                "8",
                "8",
                "EML sintéticos reservados; 4 por idioma y clase",
                "Comparación final de los tres modos",
            ],
            [
                "Validación Zenodo",
                "100",
                "23",
                "77",
                "100 textos únicos de 2.000 filas",
                "Diagnóstico secundario",
            ],
            [
                "DIFrauD",
                "1.528",
                "608",
                "920",
                "Split público histórico de phishing",
                "Diagnóstico externo con riesgo de solapamiento",
            ],
        ],
        [1200, 850, 850, 1000, 2600, 1884],
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "La separación se realiza antes del ajuste. Español conserva el test oficial y "
        "elimina del entrenamiento cualquier coincidencia exacta; inglés deduplica primero "
        "y después divide 80/20 de forma estratificada y determinista. Otros 40 casos "
        "calibran la fusión mediante cinco particiones y 16 EML distintos comprueban los "
        "tres modos con MIME y cabeceras. Zenodo y DIFrauD añaden diagnósticos externos; "
        "DIFrauD no se considera independiente por posible solapamiento de fuentes "
        "históricas con el agregado inglés (Boumber et al., 2024).",
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "evaluation/training_sources.json fija las URLs y huellas de los CSV externos. "
        "scripts/retrain_reproducible.py verifica esos archivos, reconstruye las mismas "
        "particiones, entrena y genera métricas y matrices de confusión. Los SMS españoles "
        "son una aproximación a smishing y la clase positiva inglesa también agrega spam; "
        "además, los holdouts de texto carecen de MIME. Estas limitaciones de validez de "
        "constructo impiden extrapolar las cifras a producción.",
    )
    _insert_paragraph_before(doc, anchor, "Resultados comparados", style="Heading 2")
    _insert_paragraph_before(
        doc,
        anchor,
        "Tras fijar 45 % de peso heurístico, 55 % neuronal, umbral 21 y conservación "
        "de evidencia individual a partir de 70, los holdouts y los 16 EML se procesaron sin "
        "reajustar parámetros. Accuracy mide el acierto total, precisión la fiabilidad de "
        "las alertas, recall la cobertura del phishing real y F1 su equilibrio.",
    )
    _insert_caption(doc, anchor, "Tabla 6.2: Métricas sobre los holdouts de texto.")
    _insert_table_before(
        doc,
        anchor,
        ["Idioma/modo", "N", "Accuracy", "Precisión", "Recall", "F1"],
        [
            ["ES heurístico", "209", "52,2 %", "0,0 %", "0,0 %", "0,0 %"],
            ["ES neuronal", "209", "87,1 %", "82,9 %", "92,0 %", "87,2 %"],
            ["ES combinado", "209", "89,5 %", "87,5 %", "91,0 %", "89,2 %"],
            ["EN heurístico", "16.416", "48,0 %", "92,7 %", "0,4 %", "0,9 %"],
            ["EN neuronal", "16.416", "98,3 %", "98,0 %", "98,8 %", "98,4 %"],
            ["EN combinado", "16.416", "98,4 %", "98,2 %", "98,7 %", "98,4 %"],
        ],
        [1700, 800, 1471, 1471, 1471, 1471],
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "El heurístico apenas activa señales en texto plano porque no recibe cabeceras, "
        "HTML estructurado ni autenticación. Estas filas miden principalmente el modelo "
        "léxico y la fusión; la prueba EML posterior es la comparación funcional del "
        "sistema completo.",
    )
    _insert_caption(doc, anchor, "Tabla 6.3: Métricas sobre los 16 EML reservados.")
    _insert_table_before(
        doc,
        anchor,
        ["Modo", "Accuracy", "Precisión", "Recall", "F1", "Accuracy balanceada"],
        [
            ["Heurístico", "100,0 %", "100,0 %", "100,0 %", "100,0 %", "100,0 %"],
            ["Neuronal", "81,2 %", "77,8 %", "87,5 %", "82,3 %", "81,2 %"],
            ["Combinado 45/55", "87,5 %", "80,0 %", "100,0 %", "88,9 %", "87,5 %"],
        ],
        [1800, 1280, 1280, 1160, 1160, 1704],
    )
    _insert_caption(
        doc, anchor, "Tabla 6.4: Matrices de confusión sobre los 16 EML reservados."
    )
    _insert_table_before(
        doc,
        anchor,
        ["Modo", "VN", "FP", "FN", "VP", "Lectura"],
        [
            [
                "Heurístico",
                "8",
                "0",
                "0",
                "8",
                "Clasifica correctamente los 16 escenarios",
            ],
            [
                "Neuronal",
                "6",
                "2",
                "1",
                "7",
                "Dos falsas alarmas y un phishing omitido",
            ],
            [
                "Combinado 45/55",
                "6",
                "2",
                "0",
                "8",
                "Detecta los 8 phishing y genera dos falsas alarmas",
            ],
        ],
        [1800, 700, 700, 700, 700, 3784],
    )
    _insert_paragraph_before(
        doc,
        anchor,
        "El heurístico obtiene el mejor resultado en este conjunto pequeño porque los "
        "EML contienen las cabeceras, enlaces y escenarios para los que se diseñaron las "
        "reglas. El combinado conserva los ocho verdaderos positivos y reduce el riesgo "
        "de depender solo del texto, a costa de dos falsos positivos. El neuronal genera "
        "dos falsas alarmas y omite un phishing. En los diagnósticos externos, el combinado "
        "alcanza 70,0 % de accuracy y 95,7 % de recall sobre los 100 textos únicos de "
        "Zenodo, y 89,0 % de accuracy y 93,4 % de recall sobre DIFrauD. La diferencia frente "
        "al 98,4 % del holdout inglés revela cambio de distribución y refuerza que no existe "
        "un ganador universal ni una métrica extrapolable a producción.",
    )


def _build_visual_evidence(doc: DocumentObject) -> None:
    marker = "Evidencia visual del prototipo"
    if any(p.text == marker for p in doc.paragraphs):
        _remove_between(doc, marker, "Discusión")
    anchor = _paragraph(doc, "Discusión")
    _insert_paragraph_before(doc, anchor, marker, style="Heading 3")
    _insert_paragraph_before(
        doc,
        anchor,
        "Las siguientes capturas se generaron con backend y Streamlit en procesos "
        "separados y un navegador Chromium real. La primera confirma la conexión del "
        "cliente; la segunda muestra la salida del modo combinado para el EML BEC "
        "sintético reservado, sin credenciales ni datos personales.",
    )
    _insert_picture(
        doc,
        anchor,
        ROOT / "docs" / "images" / "figura_6_1_inicio_cliente_servidor.png",
        width=6.25,
        alt_text="Pantalla inicial de Streamlit con el backend central conectado.",
    )
    _insert_caption(
        doc,
        anchor,
        "Figura 6.1: Pantalla inicial del cliente web con el backend central conectado. Fuente: elaboración propia.",
        keep_with_next=False,
    )
    _insert_picture(
        doc,
        anchor,
        ROOT / "docs" / "images" / "figura_6_2_resultado_bec_combinado.png",
        width=6.25,
        alt_text="Tarjeta de resultado combinado para un escenario BEC sintético.",
    )
    _insert_caption(
        doc,
        anchor,
        "Figura 6.2: Resultado combinado del escenario BEC sintético reservado. Fuente: elaboración propia.",
        keep_with_next=False,
    )


def _set_reference(
    paragraph: Paragraph,
    before: str,
    italic: str,
    after: str,
) -> None:
    paragraph.clear()
    paragraph.add_run(before)
    italic_run = paragraph.add_run(italic)
    italic_run.italic = True
    paragraph.add_run(after)


def _bibliography(doc: DocumentObject) -> None:
    remove_prefixes = (
        "Hoxhunt. (2025).",
        "LevelBlue. (2024).",
        "OpenAI. (2023).",
        "Proofpoint. (2025). What is phishing?",
        "Proofpoint. (2025b).",
        "Verizon. (2021).",
    )
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(remove_prefixes):
            _remove(paragraph._p)
    _replace_fragment(doc, "Proofpoint, 2025a", "Proofpoint, 2025")
    _replace_prefix(
        doc,
        "Proofpoint. (2025a).",
        "Proofpoint. (2025). The Human Factor 2025: Vol. 1 Social Engineering. "
        "https://www.proofpoint.com/us/resources/threat-reports/human-factor-social-engineering",
    )
    _replace_prefix(
        doc,
        "GreatHorn. (2021).",
        "GreatHorn. (2021). 3 phases of the phishing attack kill chain. "
        "https://web.archive.org/web/20231201201455/https://www.greathorn.com/blog/"
        "what-are-the-3-phases-of-the-phishing-attack-kill-chain/",
    )
    formatted = {
        "AlEroud, A., & Karabatis, G. (2020).": (
            "AlEroud, A., & Karabatis, G. (2020). Bypassing detection of URL-based phishing attacks using generative adversarial deep neural networks. In ",
            "Proceedings of the Sixth International Workshop on Security and Privacy Analytics (IWSPA '20)",
            " (pp. 53-60). ACM. https://doi.org/10.1145/3375708.3380315",
        ),
        "Alghenaim, M., Alkawsi, G., & Barnhart, C. R. (2025).": (
            "Alghenaim, M., Alkawsi, G., & Barnhart, C. R. (2025). The state of the art in AI-based phishing detection: A systematic literature review. In M. A. Al-Sharafi, M. Al-Emran, M. A. Mahmoud, & I. Arpaci (Eds.), ",
            "Current and Future Trends on AI Applications",
            " (Studies in Computational Intelligence, Vol. 1178, pp. 431-458). Springer. https://doi.org/10.1007/978-3-031-75091-5_23",
        ),
        "Hadnagy, C. (2018).": (
            "Hadnagy, C. (2018). ",
            "Social Engineering: The Science of Human Hacking",
            " (2nd ed.). Wiley.",
        ),
        "Hutchinson, S., Zhang, Z., & Liu, Q. (2018).": (
            "Hutchinson, S., Zhang, Z., & Liu, Q. (2018). Detecting phishing websites with random forest. In ",
            "Machine Learning and Intelligent Communications",
            " (pp. 470-479). Springer. https://doi.org/10.1007/978-3-030-00557-3_46",
        ),
        "Wang, Y., Zhu, W., Xu, H., Qin, Z., Ren, K., & Ma, W. (2023).": (
            "Wang, Y., Zhu, W., Xu, H., Qin, Z., Ren, K., & Ma, W. (2023). A large-scale pretrained deep model for phishing URL detection. In ",
            "ICASSP 2023 - 2023 IEEE International Conference on Acoustics, Speech and Signal Processing (ICASSP)",
            " (pp. 1-5). IEEE. https://doi.org/10.1109/ICASSP49357.2023.10095719",
        ),
        "Yerima, S. Y., & Alzaylaee, M. K. (2020).": (
            "Yerima, S. Y., & Alzaylaee, M. K. (2020). High accuracy phishing detection based on convolutional neural networks. In ",
            "Proceedings of the 3rd International Conference on Computer Applications & Information Security",
            " (pp. 1-6). IEEE. https://doi.org/10.1109/ICCAIS48893.2020.9096869",
        ),
    }
    for paragraph in doc.paragraphs:
        for prefix, parts in formatted.items():
            if paragraph.text.startswith(prefix):
                _set_reference(paragraph, *parts)
                break
    if not any(p.text.startswith("Boumber, D. A.") for p in doc.paragraphs):
        anchor = next(p for p in doc.paragraphs if p.text.startswith("Cialdini, R. B."))
        reference = _insert_paragraph_before(
            doc,
            anchor,
            "Boumber, D. A., Qachfar, F. Z., & Verma, R. (2024). Domain-agnostic "
            "adapter architecture for deception detection: Extensive evaluations with "
            "the DIFrauD benchmark. In Proceedings of the 2024 Joint International "
            "Conference on Computational Linguistics, Language Resources and Evaluation "
            "(LREC-COLING 2024) (pp. 5260-5274). ELRA and ICCL. "
            "https://aclanthology.org/2024.lrec-main.468/",
        )
        reference.paragraph_format.first_line_indent = Inches(-0.25)
        reference.paragraph_format.left_indent = Inches(0.25)


def apply_memory(path: Path) -> None:
    doc = Document(path)
    ai_disclosure = (
        "Durante el desarrollo se emplearon herramientas de inteligencia artificial "
        "generativa como apoyo para consultas técnicas puntuales, diagnóstico y resolución "
        "de errores, revisión de fragmentos de código, asistencia en el CSS de la interfaz "
        "y revisión de documentación. Antes de incorporar los cambios se realizaron las "
        "comprobaciones correspondientes; las decisiones técnicas, la interpretación de "
        "los resultados y la responsabilidad final corresponden al autor. Estas herramientas "
        "no se utilizaron como fuente académica: los contenidos teóricos se contrastaron y "
        "citaron mediante sus fuentes originales."
    )
    if not any(p.text == ai_disclosure for p in doc.paragraphs):
        anchor = _paragraph(
            doc,
            "La gestión del proyecto se apoyó en el control de versiones Git y en la "
            "plataforma GitHub como repositorio remoto. El desarrollo se organizó por "
            "funcionalidades y cada hito se validó mediante pruebas automatizadas antes "
            "de integrarse en la versión de trabajo, con especial atención a las funciones "
            "críticas de detección, normalización y persistencia de modelos.",
        )
        _insert_paragraph_before(doc, anchor, ai_disclosure)
    _ensure_front_matter_entries(doc)
    _build_proposal_changes(doc)
    _build_experiment(doc)
    _build_visual_evidence(doc)
    _replace_prefix(
        doc,
        "En primer lugar, se analizan las cabeceras del correo.",
        "En primer lugar, se analizan las cabeceras del correo. El sistema interpreta "
        "los resultados que los servidores de correo ya han escrito en Received-SPF, "
        "Authentication-Results y ARC-Authentication-Results, y detecta estados fail, "
        "softfail o errores de SPF, DKIM y DMARC. También revisa incoherencias entre "
        "From, Reply-To y Return-Path y si una cabecera DKIM-Signature presente está "
        "mal formada. No realiza consultas DNS, no recupera claves públicas y no "
        "verifica criptográficamente SPF, DKIM o DMARC; por tanto, es una interpretación "
        "pasiva de cabeceras, no una validación completa de autenticación.",
    )
    _replace_prefix(
        doc,
        "Para cada correo, el sistema devuelve una puntuación global",
        "Para cada correo, el sistema devuelve una puntuación global de riesgo, una "
        "clasificación y un desglose de señales. Interpreta los resultados SPF/DKIM/DMARC "
        "ya presentes en las cabeceras, sin consultas DNS ni validación criptográfica, y "
        "revisa coherencia del remitente, urgencia, URLs, IPs, punycode, acortadores, "
        "redirecciones, adjuntos e ingeniería social. Las señales activadas se acompañan "
        "de explicaciones para interpretar por qué se ha clasificado el mensaje.",
    )
    _replace_prefix(
        doc,
        "El código completo del prototipo se encuentra disponible",
        "El código completo del prototipo, las dependencias fijadas, los scripts de "
        "evaluación y las instrucciones de ejecución están disponibles en "
        "https://github.com/villarrubi/TFG.",
    )
    _replace_prefix(
        doc,
        "2024-2025: Transformers y modelos de lenguaje.",
        "2024-2025: Transformers y modelos de lenguaje. Los trabajos recientes "
        "estudian clasificación zero-shot, ajuste fino, prompt engineering y "
        "explicaciones generadas por LLM. El potencial es alto, pero la precisión, la "
        "fidelidad de las explicaciones, el coste y la privacidad no mejoran siempre de "
        "forma simultánea (Koide et al., 2024; Trad & Chehab, 2024; Kuikel et al., 2025).",
    )
    _replace_prefix(
        doc,
        "Los enfoques clásicos extraen características de URLs",
        "Los enfoques clásicos extraen características de URLs, páginas, cabeceras o "
        "texto y las proporcionan a clasificadores supervisados. Los modelos que "
        "aprenden representaciones directamente de la URL reducen parte de esa "
        "ingeniería manual (Le et al., 2018), pero su rendimiento también puede caer "
        "cuando cambia la distribución o el atacante modifica los rasgos esperados.",
    )
    conclusions = next(
        (
            p
            for p in doc.paragraphs
            if p.text == "Conclusiones" and p.style.name == "Heading 1"
        ),
        None,
    )
    if conclusions is not None:
        _set_text(conclusions, "Conclusiones generales")
    _replace_prefix(
        doc,
        "El Trabajo Fin de Grado ha alcanzado los objetivos planteados.",
        "Estas conclusiones generales valoran el trabajo completo y se diferencian de "
        "las conclusiones del estado del arte. El Trabajo Fin de Grado ha alcanzado los "
        "objetivos planteados: se ha estudiado el phishing desde su base psicológica y "
        "su evolución reciente, se ha revisado la investigación en machine learning y "
        "deep learning y se ha desarrollado un sistema cliente-servidor funcional que "
        "combina análisis heurístico explicable con un clasificador TF-IDF + MLP. "
        "Streamlit, la extensión de Gmail y el monitor envían las entradas al backend "
        "central, que procesa el correo y mantiene los modelos compartidos.",
    )
    _bibliography(doc)
    doc.save(path)


def apply_full_guide(path: Path) -> None:
    doc = Document(path)
    ai_question = "¿Has utilizado inteligencia artificial durante el desarrollo?"
    ai_answer = (
        "Sí, como herramienta de apoyo para consultas técnicas puntuales, diagnóstico de "
        "errores, revisión de fragmentos de código, CSS y revisión de documentación. Cada "
        "cambio incorporado se comprobó y las decisiones técnicas, la interpretación de "
        "resultados y la responsabilidad final son del autor. La IA no se cita como fuente "
        "académica: la parte teórica se contrastó con las referencias originales."
    )
    if not any(p.text == ai_question for p in doc.paragraphs):
        anchor = _paragraph(doc, "Demostración en directo")
        _insert_paragraph_before(
            doc, anchor, ai_question, style="Pregunta tribunal"
        )
        _insert_paragraph_before(
            doc, anchor, ai_answer, style="Respuesta tribunal"
        )
    _replace_prefix(
        doc,
        "analizador_email.py usa BytesParser",
        "analizador_email.py usa BytesParser con policy.default. Si el mensaje es "
        "multipart, recorre sus partes y conserva el primer cuerpo text/plain, el "
        "primer text/html y los nombres de los adjuntos. Si solo existe HTML, extrae "
        "el texto visible y también obtiene las anclas, sus href y el texto mostrado. "
        "Conserva las cabeceras repetidas. SPF, DKIM y DMARC se interpretan únicamente desde "
        "Received-SPF, Authentication-Results, ARC-Authentication-Results y "
        "DKIM-Signature. El prototipo no consulta DNS, no recupera claves públicas y no "
        "realiza una validación criptográfica completa.",
    )
    _replace_prefix(
        doc,
        "El repositorio separa 40 casos de calibración",
        "El protocolo reproducible usa 1.148 textos ES para entrenar y 209 para probar; "
        "el inglés usa 65.661 para entrenar y 16.416 para probar tras deduplicar y dividir "
        "el agregado con semilla 42. Otros 40 casos calibran y 16 EML finales comparan los "
        "tres modos con cabeceras. En los EML, el heurístico obtiene 100,0 % de accuracy; "
        "el combinado, 87,5 % de accuracy y 100,0 % de recall; y el neuronal, 81,2 % de "
        "accuracy. DIFrauD añade 1.528 textos con riesgo de solapamiento documentado.",
    )
    doc.save(path)


def main() -> None:
    apply_memory(ROOT / "TFG.docx")
    apply_full_guide(ROOT / "Guia_defensa_TFG.docx")
    print("Observaciones del tutor aplicadas a memoria y guía extensa.")


if __name__ == "__main__":
    main()
